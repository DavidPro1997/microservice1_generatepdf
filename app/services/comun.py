import os, requests, traceback 
from docx import Document # type: ignore
from docx.shared import Pt, Inches # type: ignore
from docx.shared import Cm, RGBColor # type: ignore
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT # type: ignore
import subprocess
import sys, platform
import base64
import re
import logging, json, shutil
import fitz #type: ignore
from PyPDF2 import PdfMerger, PdfReader, PdfWriter # type: ignore
from io import BytesIO
from app.config import Config
from PIL import Image, ImageDraw, ImageFont # type: ignore
from reportlab.pdfgen import canvas # type: ignore
from docx.oxml.ns import nsdecls,qn # type: ignore
from docx.oxml import parse_xml, OxmlElement # type: ignore
import openai #type: ignore
import app.logger_config 


class Docx:
    @staticmethod
    def reemplazar_texto_tabla_parrafo(archivo_entrada, archivo_salida, variables, estilos, alineacion="JUSTIFY"):
        try:
            doc = Document(archivo_entrada)
            for para in doc.paragraphs:
                for var, valor in variables.items():
                    if isinstance(valor, list):  # Si el valor es una lista (como paquete_incluye)
                        valor = "\n".join(str(item) if isinstance(item, dict) else item for item in valor)
                    marcador = f"[{var}]"
                    if marcador in para.text:
                        para.text = para.text.replace(marcador, str(valor))
                        for run in para.runs:
                            run.font.name = estilos["fuente"]
                            run.font.size = Pt(estilos["numero"])
                            if "color" in estilos and estilos["color"]:
                                color_hex = estilos["color"]
                                if color_hex:  # Si hay un color especificado
                                    # Convertir el color hexadecimal a RGB
                                    rgb = RGBColor(int(color_hex[1:3], 16), int(color_hex[3:5], 16), int(color_hex[5:7], 16))
                                    run.font.color.rgb = rgb
                        # Justificar el párrafo
                        if alineacion == "CENTER":
                            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        else:
                            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        

            #Recorrer las tablas del documento
            for tabla in doc.tables:
                for fila in tabla.rows:
                    for celda in fila.cells:
                        for para in celda.paragraphs:
                            for var, valor in variables.items():
                                if isinstance(valor, list):  # Si el valor es una lista
                                    valor = "\n".join(str(item) if isinstance(item, dict) else item for item in valor)
                                marcador = f"[{var}]"
                                if marcador in para.text:
                                    para.text = para.text.replace(marcador, str(valor))
                                    for run in para.runs:
                                        run.font.name = estilos["fuente"]
                                        run.font.size = Pt(estilos["numero"])
                                        if "color" in estilos and estilos["color"]:
                                            color_hex = estilos["color"]
                                            if color_hex:  # Si hay un color especificado
                                                # Convertir el color hexadecimal a RGB
                                                rgb = RGBColor(int(color_hex[1:3], 16), int(color_hex[3:5], 16), int(color_hex[5:7], 16))
                                                run.font.color.rgb = rgb
            doc.save(archivo_salida)
            return True
        except Exception as e:
            logging.error(f"Error: {e}")
            print(f"Error: {e}")  # Imprime el error si ocurre
            return False  # En caso de error, devolver False

    @staticmethod
    def reemplazar_texto_parrafos(archivo_entrada, archivo_salida, variables, estilos, alineacion="JUSTIFY"):
        try:
            # Cargar el documento
            doc = Document(archivo_entrada)

            # Reemplazar texto en los párrafos
            for para in doc.paragraphs:
                for var, valor in variables.items():
                    if isinstance(valor, list):  # Si el valor es una lista (como paquete_incluye)
                        valor = "\n".join(str(item) if isinstance(item, dict) else item for item in valor)
                    marcador = f"[{var}]"
                    if marcador in para.text:
                        # Limpiar el párrafo y reemplazar solo el marcador
                        para.clear()  # Limpiar el párrafo
                        # Agregar el texto reemplazado
                        run = para.add_run(str(valor))
                        Docx.aplicar_estilos_parrafos(run, estilos)

                        # Justificar solo el texto reemplazado
                        if alineacion == "CENTER":
                            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        else:
                            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # Guardar el documento modificado
            doc.save(archivo_salida)
            return True
        except Exception as e:
            logging.error(f"Error: {e}")
            print(f"Error: {e}")  # Imprime el error si ocurre
            return False  # En caso de error, devolver False
        
    @staticmethod
    def reemplazar_texto_tablas(archivo_entrada, archivo_salida, variables, estilos, alineacion="JUSTIFY"):
        try:
            alineaciones = {
                "LEFT": WD_PARAGRAPH_ALIGNMENT.LEFT,
                "CENTER": WD_PARAGRAPH_ALIGNMENT.CENTER,
                "RIGHT": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                "JUSTIFY": WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            }
            # Cargar el documento
            doc = Document(archivo_entrada)

            # Reemplazar texto en las tablas
            for tabla in doc.tables:
                for fila in tabla.rows:
                    for celda in fila.cells:
                        for para in celda.paragraphs:
                            for var, valor in variables.items():
                                if isinstance(valor, list):  # Si el valor es una lista
                                    valor = "\n".join(str(item) if isinstance(item, dict) else item for item in valor)
                                marcador = f"[{var}]"
                                if marcador in para.text:
                                    para.text = para.text.replace(marcador, str(valor))
                                    Docx.aplicar_estilos_tablas(para, estilos)
                                    para.alignment = alineaciones.get(alineacion.upper(), WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
                                    

            # Guardar el documento modificado
            doc.save(archivo_salida)
            return True
        except Exception as e:
            logging.error(f"Error: {e}")
            print(f"Error: {e}")  # Imprime el error si ocurre
            return False  # En caso de error, devolver False


    @staticmethod
    def reemplazar_texto_tabla_anidada(archivo_entrada, archivo_salida, variables, estilos, numero_tabla, numero_fila, numero_celda, alineacion="JUSTIFY"):
        try:
            alineaciones = {
                "LEFT": WD_PARAGRAPH_ALIGNMENT.LEFT,
                "CENTER": WD_PARAGRAPH_ALIGNMENT.CENTER,
                "RIGHT": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                "JUSTIFY": WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            }

            # Cargar el documento
            doc = Document(archivo_entrada)

            # Acceder a la primera tabla
            if not doc.tables:
                print("No hay tablas en el documento.")
                return False
            
            tabla_principal = doc.tables[numero_tabla]  # Primera tabla

            # Verificar si hay una tabla dentro de la primera celda de la primera fila
            if not tabla_principal.rows[numero_fila].cells[numero_celda].tables:
                print("No hay una tabla anidada en la primera celda de la primera fila.")
                return False
            
            tabla_anidada = tabla_principal.rows[numero_fila].cells[numero_celda].tables[0]  # Primera tabla dentro de la primera tabla

            # Reemplazar texto en la tabla anidada
            for fila in tabla_anidada.rows:
                for celda in fila.cells:
                    for para in celda.paragraphs:
                        for var, valor in variables.items():
                            marcador = f"[{var}]"
                            if marcador in para.text:
                                para.text = para.text.replace(marcador, str(valor))
                                Docx.aplicar_estilos_tablas(para, estilos)
                                para.alignment = alineaciones.get(alineacion.upper(), WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

            # Guardar el documento modificado
            doc.save(archivo_salida)
            return True

        except Exception as e:
            print(f"Error: {e}")
            return False  # En caso de error, devolver False


    @staticmethod
    def aplicar_estilos_tablas(para, estilos):
        for run in para.runs:
            run.font.name = estilos["fuente"]
            run.font.size = Pt(estilos["numero"])
            if "color" in estilos and estilos["color"]:
                color_hex = estilos["color"]
                if color_hex:  # Si hay un color especificado
                    # Convertir el color hexadecimal a RGB
                    rgb = RGBColor(int(color_hex[1:3], 16), int(color_hex[3:5], 16), int(color_hex[5:7], 16))
                    run.font.color.rgb = rgb

    @staticmethod
    def aplicar_estilos_parrafos(run, estilos):
        run.font.name = estilos["fuente"]
        run.font.size = Pt(estilos["numero"])
        if "color" in estilos and estilos["color"]:
            color_hex = estilos["color"]
            if color_hex:  # Si hay un color especificado
                # Convertir el color hexadecimal a RGB
                rgb = RGBColor(int(color_hex[1:3], 16), int(color_hex[3:5], 16), int(color_hex[5:7], 16))
                run.font.color.rgb = rgb

    @staticmethod
    def crear_tabla_rooms(archivo_entrada, archivo_salida, variable, datos, estilos):
        try:
            doc = Document(archivo_entrada)

            for para in doc.paragraphs:
                if variable in para.text:
                    para.clear()  # Eliminar el contenido del párrafo con la variable

                    for diccionario in datos:
                        # Crear una tabla para cada diccionario
                        table = doc.add_table(rows=5, cols=4)
                        table.style = 'Table Grid'
                        for row in table.rows:
                            for cell in row.cells:
                                tc = cell._element
                                tc_pr = tc.get_or_add_tcPr()
                                tc_borders = parse_xml(r'''
                                    <w:tcBorders %s>
                                        <w:top w:val="single" w:sz="4" w:space="0" w:color="E7E6E6"/>
                                        <w:left w:val="single" w:sz="4" w:space="0" w:color="E7E6E6"/>
                                        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E7E6E6"/>
                                        <w:right w:val="single" w:sz="4" w:space="0" w:color="E7E6E6"/>
                                    </w:tcBorders>''' % nsdecls('w'))
                                tc_pr.append(tc_borders)
                                # Establecer el color de fondo de la celda
                                shading_elm = parse_xml(r'<w:shd {} w:fill="E7E6E6"/>'.format(nsdecls('w')))
                                tc_pr.append(shading_elm)
                        comentarios = []
                        pares = list(diccionario.items())
                        for row_idx, row in enumerate(table.rows):
                            for col_idx, cell in enumerate(row.cells):
                                # Asignar las claves y valores según la posición
                                index = row_idx * 2 + col_idx // 2
                                if index < len(pares):
                                    clave, valor = pares[index]
                                    if clave == "rate_comments":
                                        comentarios.append(str(valor))
                                    else:
                                        if col_idx % 2 == 0:
                                            cell.text = Docx.traducir_palabras(clave)  # Clave en columna izquierda
                                            # Aplicar negrita en la columna de claves
                                            for paragraph in cell.paragraphs:
                                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                                                for run in paragraph.runs:
                                                    run.bold = True
                                        else:
                                            cell.text = str(valor)  # Valor en columna derecha

                                        # Aplicar estilos generales
                                        for paragraph in cell.paragraphs:
                                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                                            for run in paragraph.runs:
                                                run.font.name = estilos["fuente"]
                                                run.font.size = Pt(estilos["numero"])

                        # Insertar un párrafo vacío y luego la tabla
                        empty_paragraph = para.insert_paragraph_before()
                        empty_paragraph.text = ""  # Salto de línea entre tablas
                        table_element = table._element
                        empty_paragraph._element.addnext(table_element)  # Agregar la tabla después del párrafo vacío
                        if comentarios:
                            comentario_paragraph = empty_paragraph.insert_paragraph_before()
                            comentario_paragraph.text = "Comentarios: " + "; ".join(comentarios)
                            for run in comentario_paragraph.runs:
                                run.font.name = estilos["fuente"]
                                run.font.size = Pt(estilos["numero"])
                            table_element.addnext(comentario_paragraph._element)
            doc.save(archivo_salida)
            return True
        except Exception as e:
            logging.error(f"Error: {e}")
            print(f"Error: {e}")  # Imprime el error si ocurre
            return False  # En caso de error, devolver False

    @staticmethod
    def armar_tabla_vuelos(archivo_entrada, archivo_salida, variable,datos ,estilos):
        columnas = ["clase", "detalle_salida", "duracion", "detalle_destino"]
        numeroFilas = len(datos)
        ancho_columnas = [Inches(0.9), Inches(2.2), Inches(0.5), Inches(2.3)]
        try:
            doc = Document(archivo_entrada)
            for para in doc.paragraphs:
                if variable in para.text:
                    para.clear()
                    table = doc.add_table(rows=numeroFilas, cols=4)
                    table.style = 'Plain Table 2'

                    # Aplicar ancho a las columnas
                    for i, column in enumerate(table.columns):
                        column.width = ancho_columnas[i]
                    
                    # Primero agregamos las filas de datos a la tabla
                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            if i < len(datos):  # Validar que exista la fila
                                clave = columnas[j]  # Obtener la clave correspondiente a la columna
                                if clave in datos[i]:  # Validar que la clave exista en el diccionario
                                    valor = datos[i][clave]
                                    if isinstance(valor, list):
                                        valor = "\n".join(valor)
                                    cell.text = valor
                                    for paragraph in cell.paragraphs:
                                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    for run in paragraph.runs:
                                        run.bold = False  # Esto quita la negrita

                       # Verificar si escala tiene valor y agregar fila adicional
                        escala = datos[i].get("escala", "")
                        if escala:  # Si escala no está vacío
                            new_row = table.add_row()  # Agregar una nueva fila
                            new_row._tr.addnext(row._tr)  # Insertar después de la fila actual
                            merged_cell = new_row.cells[0]  # Seleccionar la primera celda de la fila
                            merged_cell.merge(new_row.cells[-1])  # Unir todas las columnas en una sola
                            merged_cell.text = escala  # Colocar el texto de escala
                            for paragraph in merged_cell.paragraphs:
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            for run in paragraph.runs:
                                run.bold = False

                    table_element = table._element
                    para._element.addnext(table_element)
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = estilos["fuente"]
                                    run.font.size = Pt(estilos["numero"])
            doc.save(archivo_salida)
            return True
        except Exception as e:
            logging.error(f"Error: {e}")
            print(f"Error: {e}")  # Imprime el error si ocurre
            return False  # En caso de error, devolver False

    @staticmethod
    def imagen_en_docx(image_path, docx_path, key, ancho_en_pt=None, alto_en_pt=None, alineacion="CENTER"):
        try:
            doc = Document(docx_path)

            image = Image.open(image_path)

            # Calcular el ancho en proporción al alto especificado
            ancho_original, alto_original = image.size
            
            if ancho_en_pt is not None and alto_en_pt is None:
                # Si solo se proporciona el ancho, calcular el alto manteniendo la proporción
                proporción = ancho_en_pt / ancho_original
                alto_en_pt = int(alto_original * proporción)
            elif alto_en_pt is not None and ancho_en_pt is None:
                # Si solo se proporciona el alto, calcular el ancho manteniendo la proporción
                proporción = alto_en_pt / alto_original
                ancho_en_pt = int(ancho_original * proporción)
            elif alto_en_pt is None and ancho_en_pt is None:
                alto_en_pt = alto_original
                ancho_en_pt = ancho_original


            if ancho_en_pt > 400:
                ancho_en_pt = 400
                proporción = ancho_en_pt / ancho_original
                alto_en_pt = int(alto_original * proporción)

            extension = os.path.splitext(image_path)[1].lower()
            if extension == '.png':
                formato_imagen = 'PNG'
            else:
                formato_imagen = 'JPEG'

            # Guardar la imagen en un buffer de memoria
            image_stream = BytesIO()
            image.save(image_stream, format=formato_imagen)  # Guardar como JPEG
            image_stream.seek(0)

            # Bandera para verificar si se encontró la clave
            found = False

            # Buscar la clave en las celdas de la tabla
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if key in cell.text:
                            found = True
                            
                            # Limpiar el contenido de la celda
                            for paragraph in cell.paragraphs:
                                paragraph.clear()  # Limpiar el contenido del párrafo

                            # Insertar la imagen en la celda y centrarla
                            run = cell.paragraphs[0].add_run()
                            run.add_picture(image_stream, width=Pt(ancho_en_pt), height=Pt(alto_en_pt))
                            if alineacion == "CENTER":
                                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar el párrafo
                            elif alineacion == "JUSTIFY":
                                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            elif alineacion == "RIGTH":
                                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGTH  # Centrar el párrafo
                            elif alineacion == "LEFT":
                                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Centrar el párrafo
                            break
            
            # Buscar la clave en los párrafos normales si no está en una tabla
            if not found:
                for paragraph in doc.paragraphs:
                    if key in paragraph.text:
                        # Limpiar el contenido del párrafo
                        paragraph.clear()  # Limpiar el contenido del párrafo

                        # Crear un nuevo párrafo para la imagen
                        # new_paragraph = paragraph.insert_paragraph_before()
                        new_paragraph = paragraph.insert_paragraph_before()
                        if alineacion == "CENTER":
                            new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar el párrafo
                        elif alineacion == "JUSTIFY":
                            new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        elif alineacion == "RIGHT":
                            new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Centrar el párrafo
                        elif alineacion == "LEFT":
                            new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Centrar el párrafo
                        # new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar el párrafo
                        run = new_paragraph.add_run()
                        # run = cell.paragraphs[0].add_run()
                        run.add_picture(image_stream, width=Pt(ancho_en_pt), height=Pt(alto_en_pt))
                        break

            # Guardar el documento actualizado
            doc.save(docx_path)
            return True

        except Exception as e:
            print(f"Error: {e}")
            traceback.print_exc()  # Imprime la traza del error
            return False

    @staticmethod
    def eliminar_filas_docx(ruta_entrada, ruta_salida, filas_a_eliminar, numero_tabla, numero_fila = None, numero_celda = None, numero_tabla_anidada = None):
        try:
            doc = Document(ruta_entrada)
            if numero_fila is not None and numero_celda is not None and numero_tabla_anidada is not None:
                tabla  = doc.tables[numero_tabla].rows[numero_fila].cells[numero_celda].tables[numero_tabla_anidada]
            else:
                tabla = doc.tables[numero_tabla]
            num_filas = len(tabla.rows)
            filas_a_eliminar_validas = [fila_idx for fila_idx in filas_a_eliminar if fila_idx < num_filas]
            for fila_idx in sorted(filas_a_eliminar_validas, reverse=True):
                fila = tabla.rows[fila_idx]
                tbl = tabla._element
                tbl.remove(fila._element)  # Elimina la fila del XML de la tabla
            
            doc.save(ruta_salida)
            return True
        
        except Exception as e:
            print(f"Error: {e}")
            return False

    @staticmethod
    def aplicar_estilos_especificos(archivo_entrada, archivo_salida):
        textos_objetivo = {
            'COMPAÑÍA TURISTICA "MARKETING VIP S.A" COMTUMARK',
            'EL CLIENTE',
            'SEGUNDA',
            'TERCERA'
            'LA OPERADORA TURISTICA',
            'Lugar y fecha'
        }
        try:
            doc = Document(archivo_entrada)
            for para in doc.paragraphs:
                nuevo_runs = []
                for run in para.runs:
                    texto = run.text
                    i = 0
                    while i < len(texto):
                        texto_encontrado = False
                        for palabra in textos_objetivo:
                            if texto[i:i+len(palabra)] == palabra:
                                # Crear un nuevo run con negrita para el texto objetivo
                                nuevo_run = para.add_run(palabra)
                                nuevo_run.bold = True
                                nuevo_run.font.name = 'Helvetica'
                                nuevo_run.font.size = Pt(10)
                                nuevo_runs.append(nuevo_run)
                                i += len(palabra)
                                texto_encontrado = True
                                break
                        if not texto_encontrado:
                            # Crear un run normal para el texto que no coincide
                            nuevo_run = para.add_run(texto[i])
                            nuevo_run.font.name = run.font.name or 'Helvetica'
                            nuevo_run.font.size = run.font.size or Pt(10)
                            nuevo_run.bold = run.bold
                            nuevo_runs.append(nuevo_run)
                            i += 1
                    run._element.getparent().remove(run._element)
            doc.save(archivo_salida)
            return True
        except Exception as e:
            logging.error(f"Error: {e}")
            print(f"Error: {e}")  # Imprime el error si ocurre
            return False  # En caso de error, devolver False

    # @staticmethod
    # def agregar_imagen_docx(ruta_imagen, ruta_plantilla, ruta_salida):
    #     try:
    #         # Cargar la plantilla del documento
    #         doc = Document(ruta_plantilla)
            
    #         # Abrir la imagen y obtener sus dimensiones
    #         with Image.open(ruta_imagen) as img:
    #             width, height = img.size
                
    #             # Ajustar dimensiones de la imagen
    #             if height > width:  # Si la imagen es más alta que ancha
    #                 alto = Cm(19)
    #                 ancho = (width / height) * alto
    #             else:  # Si la imagen es más ancha que alta
    #                 ancho = Cm(13)
    #                 alto = (height / width) * ancho
                
    #             # Insertar la imagen al principio del documento
    #             paragraph = doc.add_paragraph()  # Crear un nuevo párrafo al inicio
    #             run = paragraph.add_run()  # Crear un "run" en el párrafo
    #             run.add_picture(ruta_imagen, width=ancho, height=alto)  # Insertar la imagen con tamaño ajustado
    #             paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar el párrafo
            
    #         # Guardar el documento con la imagen añadida al principio
    #         doc.save(ruta_salida)
    #         return True
    #     except Exception as e:
    #         logging.error(f"Error: {e}")
    #         print(f"Error: {e}")  # Imprime el error si ocurre
    #         return False  # En caso de error, devolver False
      
    @staticmethod
    def traducir_palabras(palabra):
        if palabra == "room_name":
            return "Nombre habitación:"
        elif palabra == "acomodation":
            return "Acomodation Type:"
        elif palabra == "name_pax":
            return "Pasajero:"
        elif palabra == "adults":
            return "Adultos:"
        elif palabra == "children":
            return "Niños"
        elif palabra == "age_children":
            return "Edad niños:"
        elif palabra == "board_basis":
            return "Board Base:"
        elif palabra == "room_number":
            return "Número habitaciones:"
        elif palabra == "rate_comments":
            return "Comentario de tarifa:"
        elif palabra == "Forma de pago":
            return "Comentario de tarifa:"
        else:
            return palabra

    @staticmethod
    def reemplazar_con_hipervinculo(archivo_entrada, archivo_salida, variable, url, texto, estilos, alineacion="JUSTIFY"):
        """
        Reemplaza la palabra clave en un documento Word con un hipervínculo formateado,
        buscando en párrafos normales y dentro de tablas.

        :param archivo_entrada: Ruta del archivo Word de entrada.
        :param archivo_salida: Ruta del archivo Word de salida.
        :param variable: Palabra clave a reemplazar.
        :param url: URL del hipervínculo.
        :param texto: Texto visible del enlace.
        :param estilos: Diccionario con "fuente", "numero" (tamaño), "color" (hex).
        :param alineacion: Alineación del párrafo ("JUSTIFY", "CENTER", "LEFT", "RIGHT").
        :return: True si tuvo éxito, False si ocurrió un error.
        """
        try:
            doc = Document(archivo_entrada)
            reemplazado = False

            def procesar_parrafo(para):
                """ Busca la variable en un párrafo y la reemplaza con el hipervínculo """
                if variable in para.text:
                    # Eliminar el texto original
                    para.clear()

                    # Aplicar alineación
                    alineaciones = {
                        "JUSTIFY": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                        "CENTER": WD_PARAGRAPH_ALIGNMENT.CENTER,
                        "LEFT": WD_PARAGRAPH_ALIGNMENT.LEFT,
                        "RIGHT": WD_PARAGRAPH_ALIGNMENT.RIGHT
                    }
                    para.alignment = alineaciones.get(alineacion.upper(), WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

                    # Crear el hipervínculo
                    hyperlink = OxmlElement("w:hyperlink")

                    run = OxmlElement("w:r")
                    rPr = OxmlElement("w:rPr")

                    # Aplicar fuente y tamaño
                    if "fuente" in estilos:
                        rFonts = OxmlElement("w:rFonts")
                        rFonts.set(qn("w:ascii"), estilos["fuente"])
                        rPr.append(rFonts)
                    if "numero" in estilos:
                        sz = OxmlElement("w:sz")
                        sz.set(qn("w:val"), str(int(estilos["numero"]) * 2))  # Word usa media-puntos
                        rPr.append(sz)

                    # Aplicar color
                    if "color" in estilos:
                        color = OxmlElement("w:color")
                        color.set(qn("w:val"), estilos["color"].replace("#", ""))
                        rPr.append(color)

                    # Aplicar subrayado
                    u = OxmlElement("w:u")
                    u.set(qn("w:val"), "single")  # "single" para subrayado simple
                    rPr.append(u)

                    run.append(rPr)

                    text_elem = OxmlElement("w:t")
                    text_elem.text = texto
                    run.append(text_elem)
                    hyperlink.append(run)

                    # Relacionar con la URL
                    part = para._parent.part
                    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
                    hyperlink.set(qn("r:id"), r_id)

                    para._element.append(hyperlink)
                    return True

                return False

            # Buscar en párrafos normales
            for para in doc.paragraphs:
                if procesar_parrafo(para):
                    reemplazado = True

            # Buscar en tablas
            for tabla in doc.tables:
                for fila in tabla.rows:
                    for celda in fila.cells:
                        for para in celda.paragraphs:
                            if procesar_parrafo(para):
                                reemplazado = True

            if reemplazado:
                doc.save(archivo_salida)
                return True
            else:
                return False

        except Exception as e:
            print(f"Error al procesar el documento: {e}")
            return False




class Imagen:
    @staticmethod
    def guardar_imagen_base64(imagen_base64, ruta_salida):
        try:
            match = re.match(r"data:image/(\w+);base64,(.*)", imagen_base64)
            if not match:
                return False
            imagen_base64_data = match.group(2)
            imagen_bytes = base64.b64decode(imagen_base64_data)
            ruta_imagen = f"{ruta_salida}.jpeg"
            with open(ruta_imagen, "wb") as f:
                f.write(imagen_bytes)
                return ruta_imagen
        except Exception as e:
            logging.error(f"Error: {e}")
            print(f"Error al guardar la imagen: {e}")
            return False

    @staticmethod
    def download_image(image_url: str, save_path: str):
        try:
            # Realizar la solicitud HTTP para obtener la imagen
            response = requests.get(image_url)
            
            # Verificar que la solicitud fue exitosa
            response.raise_for_status()
            
            # Abrir la imagen desde los datos descargados
            image = Image.open(BytesIO(response.content))
            
            # Guardar la imagen en la ruta especificada sin modificar su tamaño
            image.save(save_path)
            
            return True
        
        except requests.exceptions.RequestException as e:
            # Manejar errores de la solicitud HTTP
            print(f"Error al descargar la imagen: {e}")
            return False
        except Exception as e:
            # Manejar otros posibles errores
            print(f"Error: {e}")
            return False
        
    @staticmethod
    def convertir_imagen_a_base64(ruta_imagen):
        try:
            with open(ruta_imagen, "rb") as imagen:
                datos_imagen = imagen.read()
                base64_imagen = base64.b64encode(datos_imagen).decode("utf-8")
            return base64_imagen
        except Exception as e:
            print(f"Ocurrió un error: {e}")
            logging.error(f"Ocurrió un error: {e}")
            return False

    @staticmethod
    def colocar_texto_a_imagen(texto, coordenadas, ruta_imagen, ruta_salida, fuente="arial.ttf", tamano=40, color="black", negrita=False):        
        try:
            sistema_operativo = platform.system()
            # Configurar la ruta de la fuente según el sistema operativo
            if sistema_operativo == "Windows":
                ruta_fuente = f"C:/Windows/Fonts/{fuente}"  # Ruta típica en Windows
            elif sistema_operativo == "Linux":
                ruta_fuente = f"/usr/share/fonts/truetype/dejavu/{fuente}"  # Fuente común en Linux
            else:
                logging.error("Sistema operativo no soportado para fuentes predeterminadas.")
                raise OSError("Sistema operativo no soportado para fuentes predeterminadas.")
            
            # Verificar que la fuente exista
            if not os.path.exists(ruta_fuente):
                logging.error(f"La fuente no se encuentra en {ruta_fuente}")
                raise FileNotFoundError(f"La fuente no se encuentra en {ruta_fuente}")

            imagen = Image.open(ruta_imagen)
            draw = ImageDraw.Draw(imagen)

            # Cargar la fuente con el tamaño
            if negrita:
                # Verificar si existe una fuente en negrita y cargarla
                ruta_fuente_negrita = ruta_fuente.replace(".ttf", "-Bold.ttf")
                if os.path.exists(ruta_fuente_negrita):
                    fuente = ImageFont.truetype(ruta_fuente_negrita, tamano)
                else:
                    # Si no existe la fuente en negrita, solo usar la fuente normal
                    fuente = ImageFont.truetype(ruta_fuente, tamano)
            else:
                fuente = ImageFont.truetype(ruta_fuente, tamano)

            # Dibujar el texto en la imagen con el color especificado
            draw.text(coordenadas, texto, fill=color, font=fuente)

            # Guardar la imagen con el texto añadido
            imagen.save(ruta_salida)
            return True
        except Exception as e:
            print(f"Ocurrió un error: {e}") 
            logging.error(f"Ocurrió un error: al colocar texto {e}")
            return False

    @staticmethod
    def colocar_imagen_pequena(imagen_pequena, coordenadas, ruta_imagen, ruta_salida, ancho_en_pt=None, alto_en_pt=None, rotar=0):
        try:
            # Cargar la imagen de fondo
            imagen_grande = Image.open(ruta_imagen)
            
            # Crear una nueva imagen blanca del mismo tamaño que la imagen de fondo
            fondo_blanco = Image.new("RGB", imagen_grande.size, (255, 255, 255))
            
            # Pegar la imagen grande en el fondo blanco
            fondo_blanco.paste(imagen_grande, (0, 0))
            
            # Cargar la imagen pequeña
            imagen_pequena = Image.open(imagen_pequena)

            # Calcular el ancho en proporción al alto especificado
            ancho_original, alto_original = imagen_pequena.size
            
            if ancho_en_pt is not None and alto_en_pt is None:
                # Si solo se proporciona el ancho, calcular el alto manteniendo la proporción
                proporción = ancho_en_pt / ancho_original
                alto_en_pt = int(alto_original * proporción)
            elif alto_en_pt is not None and ancho_en_pt is None:
                # Si solo se proporciona el alto, calcular el ancho manteniendo la proporción
                proporción = alto_en_pt / alto_original
                ancho_en_pt = int(ancho_original * proporción)
            elif alto_en_pt is None and ancho_en_pt is None:
                ancho_en_pt = ancho_original
                alto_en_pt = alto_original
            # else:
            #     return False

            
            # Redimensionar la imagen pequeña al tamaño especificado
            imagen_pequena = imagen_pequena.resize((ancho_en_pt, alto_en_pt), Image.LANCZOS)  # Cambiado de ANTIALIAS a LANCZOS

            if rotar != 0:
                imagen_pequena = imagen_pequena.convert("RGBA")
                imagen_pequena = imagen_pequena.rotate(rotar, expand=True)
                # Create a transparent background
                fff = Image.new("RGBA", imagen_pequena.size, (255, 255, 255, 0))
                imagen_pequena = Image.alpha_composite(fff, imagen_pequena)
            
            fondo_blanco.paste(imagen_pequena, coordenadas, 
                            imagen_pequena.convert("RGBA").getchannel("A") if imagen_pequena.mode == 'RGBA' else None)
            
            # Guardar la imagen resultante
            fondo_blanco.save(ruta_salida)  
            return True          
        except Exception as e:
            print(f"Ocurrió un error: {e}") 
            logging.error(f"Ocurrió un error al colocar imagenes: {e}")  
            return False

    @staticmethod
    def resize_and_crop(image_path, width_pt=None, height_pt=None, output_path="output.png"):
        try:
            with Image.open(image_path) as img:
                orig_width, orig_height = img.size

                if width_pt:
                    new_height = int((width_pt / orig_width) * orig_height)
                    new_size = (width_pt, new_height)
                elif height_pt:
                    new_width = int((height_pt / orig_height) * orig_width)
                    new_size = (new_width, height_pt)
                else:
                    return False  # Si no se pasa width_pt ni height_pt, retorna False

                resized_img = img.resize(new_size, Image.LANCZOS)
                resized_img.save(output_path)
                return True  # Operación exitosa

        except Exception as e:
            print(f"Error: {e}")  # Opcional: Imprimir el error para depuración
            return False  # Retorna False si ocurre un error


    @staticmethod
    def resize_image_for_pdf(image_path, output_path, target_width_px, target_height_px, dpi=96):
        try:
            image = Image.open(image_path)
            
            # Redimensionar manteniendo proporciones
            image.thumbnail((target_width_px, target_height_px), Image.LANCZOS)

            # Guardar con la resolución correcta
            image.save(output_path, dpi=(dpi, dpi))

            return True
        except Exception as e:
            print(f"Error: {e}")
            return False



class Archivos:
    @staticmethod
    def archivo_a_base64(ruta_archivo):
        try:
            with open(ruta_archivo, "rb") as archivo:
                contenido = archivo.read()
                contenido_base64 = base64.b64encode(contenido).decode('utf-8')
                return contenido_base64
        except FileNotFoundError:
            logging.error(f"El archivo {ruta_archivo} no se encuentra.")
            print(f"El archivo {ruta_archivo} no se encuentra.")
            return False
        except Exception as e:
            logging.error(f"Error al convertir el PDF a base64: {e}")
            print(f"Error al convertir el PDF a base64: {e}")
            return False

    @staticmethod
    def eliminar_documentos(rutas_documentos):
        for ruta in rutas_documentos:
            try:
                # Verificar si el archivo existe
                if os.path.exists(ruta):
                    os.remove(ruta)  # Eliminar el archivo
                else:
                    logging.error(f"El archivo {ruta} no existe.")
                    print(f"El archivo {ruta} no existe.")
            except Exception as e:
                logging.error(f"Error al intentar eliminar el archivo {ruta}: {e}")
                print(f"Error al intentar eliminar el archivo {ruta}: {e}")
                return False
        return True

    @staticmethod
    def guardar_archivo_base64(ruta_guardar, base64_string):
        try:
            contenido_binario = base64.b64decode(base64_string)
            archivo_docx = BytesIO(contenido_binario)
            doc = Document(archivo_docx)
            directorio = os.path.dirname(ruta_guardar)
            if not os.path.exists(directorio):
                os.makedirs(directorio)
            doc.save(ruta_guardar)
            return True
        except Exception as e:
            logging.error(f"Error al guardar el archivo: {e}")
            print(f"Error al guardar el archivo: {e}")
            return False

    @staticmethod
    def truncar_texto(texto, num_palabras):
        palabras = texto.split()
        if len(palabras) > num_palabras:
            return " ".join(palabras[:num_palabras]) + "..."
        return texto

    @staticmethod
    def eliminar_contenido_directorio(ruta_directorio):
        try:
            # Verificar si el directorio existe
            if os.path.exists(ruta_directorio):
                # Iterar sobre los elementos en el directorio
                for item in os.listdir(ruta_directorio):
                    item_path = os.path.join(ruta_directorio, item)
                    
                    # Si el item es un archivo, lo eliminamos
                    if os.path.isfile(item_path):
                        os.remove(item_path)
                    # Si el item es un directorio, lo eliminamos recursivamente
                    elif os.path.isdir(item_path):
                        shutil.rmtree(item_path)
                
                return True
            else:
                print(f"El directorio {ruta_directorio} no existe.")
                return False
        except Exception as e:
            logging.error(f"Error al eliminar el contenido del directorio: {e}")
            print(f"Error al eliminar el contenido del directorio: {e}")
            return False




class Pdf:
    @staticmethod
    def unir_pdfs(rutas, ruta_resultado):
        try:
            merger = PdfMerger()
            for ruta in rutas:
                merger.append(ruta)
            merger.write(ruta_resultado)
            merger.close()
            return True
        except Exception as e:
            logging.error(f"Error al combinar los PDFs: {e}")
            print(f"Error al combinar los PDFs: {e}")
            return False

    @staticmethod
    def convertir_docx_a_pdf(archivo_entrada, archivo_salida):
        try:
            if not os.path.exists(archivo_entrada):
                raise FileNotFoundError(f"El archivo {archivo_entrada} no se encuentra.")
            if sys.platform == "win32":  # Windows
                libreoffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
            elif sys.platform == "linux" or sys.platform == "linux2":  # Linux
                libreoffice_path = "/usr/bin/libreoffice"  # Asegúrate de que esté en esa ubicación
            else:
                raise EnvironmentError("Sistema operativo no soportado")
            subprocess.run([libreoffice_path, '--headless', '--convert-to', 'pdf', archivo_entrada, '--outdir', archivo_salida])
            nombre_pdf = os.path.splitext(os.path.basename(archivo_entrada))[0] + ".pdf"
            ruta_pdf_salida = os.path.join(archivo_salida, nombre_pdf)
            if not os.path.exists(ruta_pdf_salida):
                logging.error(f"No se pudo crear el archivo PDF en {ruta_pdf_salida}.")
                return False
            return ruta_pdf_salida
        except Exception as e:
            logging.error(f"Error: {e}")
            print(f"Error: {e}")  # Imprime el error si ocurre
            return False  # En caso de error, devolver False

    @staticmethod
    def imagen_a_pdf(ruta_imagen, ruta_pdf_salida):
        try:
            # Abrir la imagen con Pillow
            imagen = Image.open(ruta_imagen)
            
            # Obtener el tamaño de la imagen
            ancho, alto = imagen.size

            # Crear un archivo PDF con ReportLab
            c = canvas.Canvas(ruta_pdf_salida)
            
            # Configurar el tamaño de la página en el PDF igual al tamaño de la imagen
            c.setPageSize((ancho, alto))
            
            # Insertar la imagen en el PDF
            c.drawImage(ruta_imagen, 0, 0, width=ancho, height=alto)

            # Guardar el PDF
            c.save()

            return True

        except Exception as e:
            print(f"Error: {e}")
            return False

    @staticmethod
    def contar_paginas(pdf_path):
        try:
            with open(pdf_path, "rb") as file:
                pdf = PdfReader(file)
                return len(pdf.pages)  # Devuelve el número de páginas
        except Exception as e:
            print(f"Error: {e}")
            return False

    @staticmethod
    def eliminar_pagina(pdf_path, pagina, output_path="output.pdf"):
        try:
            reader = PdfReader(pdf_path)
            writer = PdfWriter()

            if pagina < 1 or pagina > len(reader.pages):
                print("Número de página fuera de rango.")
                return False

            # Copiar todas las páginas excepto la que queremos eliminar
            for i in range(len(reader.pages)):
                if i != pagina - 1:  # `PdfReader` indexa desde 0
                    writer.add_page(reader.pages[i])

            # Guardar el nuevo PDF sin la página eliminada
            with open(output_path, "wb") as output_file:
                writer.write(output_file)

            return True  # Éxito
        except Exception as e:
            print(f"Error: {e}")
            return False  # Fallo

    @staticmethod
    def editar_pdf(input_url, output_url, texto, x, y, estilo):
        try:
            # Cargar el PDF
            doc = fitz.open(input_url)
            pagina = doc[0]  # Edita la primera página (puedes cambiarlo)

            # Extraer el estilo
            tipografia, tamano_fuente, color = estilo
            color_rgb = tuple(c / 255 for c in color) 
            
            # Insertar el texto
            pagina.insert_text((x, y), texto, fontsize=tamano_fuente, fontname=tipografia, color=color_rgb)

            # Guardar los cambios
            doc.save(output_url)
            doc.close()
            return True
        except Exception as e:
            print(f"Error al editar el PDF: {e}")
            return False    


    @staticmethod
    def centrar_texto(texto, fontname, fontsize, rango=(100, 440)):
        """ Centra el texto en el rango especificado considerando su longitud """
        # Si el texto tiene más de 40 caracteres, acortarlo (solo la primera palabra)
        if len(texto) > 40:
            palabras = texto.split()
            palabras[0] = palabras[0][0] + "."  # Reemplazar la primera palabra con la inicial
            texto = " ".join(palabras)

        # Crear la fuente
        doc = fitz.open()  # Abrir un documento temporal solo para usar la fuente
        fuente = fitz.Font(fontname)
        
        # Calcular el ancho del texto
        ancho_texto = fuente.text_length(texto, fontsize)
        
        # Calcular el espacio disponible
        espacio_disponible = rango[1] - rango[0]
        
        # Centrar el texto: la posición x es la mitad del espacio disponible menos la mitad del ancho del texto
        x = rango[0] + (espacio_disponible - ancho_texto) / 2
        
        return x, texto



class Api:
    @staticmethod
    def llamar_api_get(url, params=None, headers=None):
        """Realiza una solicitud GET a la API y devuelve la respuesta en JSON."""
        try:
            response = requests.get(url, params=params, headers=headers)
            response.raise_for_status()  # Lanza un error si la respuesta tiene un código de error
            return response.json()  # Retorna el JSON de la respuesta
        except requests.exceptions.RequestException as e:
            print(f"Error en la solicitud GET: {e}")
            return None

    @staticmethod
    def llamar_api_post(url, data=None, headers=None):
        """Realiza una solicitud POST a la API con los datos proporcionados y devuelve la respuesta en JSON."""
        try:
            response = requests.post(url, json=data, headers=headers)
            response.raise_for_status()
            return response.json()
        except requests.exceptions.RequestException as e:
            print(f"Error en la solicitud POST: {e}")
            return None

    @staticmethod
    def open_ai(rol, mensaje):
        client = openai.OpenAI(api_key=Config.key_open_ai)
        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",  # O usa "gpt-3.5-turbo" si prefieres
                messages=[
                    {"role": "system", "content": rol},
                    {"role": "user", "content": mensaje}
                ],
                response_format={"type": "json_object"} 
            )
            # Imprimir la respuesta del asistente
            respuesta = response.choices[0].message.content
            return json.loads(respuesta)
        except openai.OpenAIError as e:
            print(f"Error en la API de OpenAI: {e}")
            return ""

        except Exception as e:
            print(f"Error inesperado: {e}")
            return ""