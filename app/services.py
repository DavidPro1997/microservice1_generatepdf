import os, requests, traceback, shutil
from docx import Document # type: ignore
from docx.shared import Pt, Inches # type: ignore
from docx.shared import Cm, RGBColor # type: ignore
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT # type: ignore
import subprocess
import sys, platform
import base64
from PIL import Image # type: ignore
import re, mimetypes
import logging
from PyPDF2 import PdfMerger # type: ignore
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont # type: ignore
from docx.oxml.ns import nsdecls # type: ignore
from docx.oxml import parse_xml # type: ignore
from app import app
from datetime import datetime
import uuid
import locale
locale.setlocale(locale.LC_TIME, "es_ES.utf8")


logging.basicConfig(
    filename = os.path.abspath("logs/output.log"), 
    level=logging.DEBUG,  # Define el nivel de los logs (INFO, DEBUG, etc.)
    format='%(asctime)s - %(levelname)s - %(message)s'
)
app.logger.setLevel(logging.DEBUG)
# Manejador global de errores
@app.errorhandler(Exception)
def handle_exception(e):
    app.logger.error("Se produjo un error: %s", e, exc_info=True)
    return "Ocurrió un error en el servidor.", 500

class Switch:
    @staticmethod
    def verificar_tipo_doc(data):
        unique_id = str(uuid.uuid4())[:8]
        if data["tipo"] == "contrato":
            logging.info("Realizando servicio de creacion de pdf")
            return Contrato.generar_contrato(data)
        elif data["tipo"] == "adendum":
            logging.info("Realizando servicio de creacion de pdf")
            return Adendum.generar_adendum(data)
        elif data["tipo"] == "cotizar_vuelo":
            logging.info("Realizando servicio de creacion de pdf")
            return Cotizador.cotizar_vuelos(data, unique_id)
        elif data["tipo"] == "cotizar_vuelo_imagen":
            logging.info("Realizando servicio de creacion de imagen")
            return Imagen.cotizar_vuelos(data)
        elif data["tipo"] == "voucher_hotel":
            logging.info("Realizando servicio de creacion de voucher")
            return Hotel.generar_voucher(data)
        elif data["tipo"] == "cotizador_general":
            logging.info("Realizando servicio de creacion de cotizacion completa")
            return Cotizador.cotizar_completo(data, unique_id)
        else:
            return {"estado": False, "mensaje": "No se reconoce el tipo de archivo"}

    @staticmethod
    def verificar_tipo_doc_descarga(id):
        logging.info("Realizacion de servicio de descarga de plantilla")
        # 0 = contratos, 1 = adendum, 2 = declaraciones
        if id == 0:
            ruta = os.path.abspath("plantilla/plantilla_contratos.docx")
        elif id == 1:
            ruta = os.path.abspath("plantilla/plantilla_adendum.docx")
        elif id == 2:
            ruta = os.path.abspath("plantilla/plantilla_declaraciones.docx")
        else:
            return {"estado": False, "mensaje": "No se reconoce el tipo de archivo"}
        return Descarga.descargar_documento(ruta)

    @staticmethod
    def verificar_tipo_doc_plantilla(data,id):
        logging.info("Realizacion de servicio de actualizacion de plantilla")
        if data["archivo"]:
            # 0 = contratos, 1 = adendum, 2 = declaraciones
            if id == 0:
                ruta = os.path.abspath("plantilla/plantilla_contratos.docx")
            elif id == 1:
                ruta = os.path.abspath("plantilla/plantilla_adendum.docx")
            elif id == 2:
                ruta = os.path.abspath("plantilla/plantilla_declaraciones.docx")
            else:
                return {"estado": False, "mensaje": "No se reconoce el tipo de archivo"}
            log_guardar = Guardar.guardar_archivo(ruta,data["archivo"])
            if log_guardar:
                return {"estado": True, "mensaje": "Se ha modificado la plantilla correctamente"}
            else:
                return {"estado": False, "mensaje": "No se ha podido guardar el archivo"}
        else:
           return {"estado": False, "mensaje": "No se envio ningun archivo"} 


class Cotizador:
    @staticmethod
    def cotizar_completo(data, idUnico):
        docs_eliminar = []
        incluye = []
        if "hotel" in data and data["hotel"]:
            ciudad = "\n".join(item["city"] for item in data["hotel"])
            # opcion1
            if "vuelo" in data and data["vuelo"]:
                ticket = "Si incluye ticket aéreo"
                portada = Cotizador.generarPDFPortada(ciudad, idUnico)
                if portada["estado"]:
                    ruta_portada = portada["ruta"]
                    docs_eliminar.append(ruta_portada)
                log_paquete = Hotel.generar_pdf_paquete(data["hotel"], data["actividades"], ticket, idUnico)
                if log_paquete["estado"]:
                    ruta_paquete = log_paquete["ruta"]
                    docs_eliminar.append(ruta_paquete)
                vuelos = Cotizador.cotizar_vuelos(data["vuelo"], idUnico)
                if vuelos["estado"]:
                    incluye.append("tickets aéreos")
                    ruta_vuelos = vuelos["ruta"]
                    docs_eliminar.append(ruta_vuelos)
                log_hotel = Hotel.generar_pdf_hotel(data["hotel"], idUnico)
                if log_hotel["estado"]:
                    incluye.append("hospedaje")
                    ruta_hotel = log_hotel["ruta"]
                    docs_eliminar.append(ruta_hotel)
                if "actividades" in data and data["actividades"]:
                    log_actividades = Actividad.generarPdfActividades(data["actividades"], idUnico)
                    if log_actividades["estado"]:
                        incluye.append("actividades")
                        ruta_actividades = log_actividades["ruta"]
                        docs_eliminar.append(ruta_actividades)
                if "costos" in data and data["costos"]:
                    costos = Costos.generarPdfCostos(data["costos"], incluye, idUnico)
                    if costos["estado"]:
                        ruta_costos = costos["ruta"]
                        docs_eliminar.append(ruta_costos)
                
            # opcion2
            else:
                ticket = "No incluye ticket aéreo"
                portada = Cotizador.generarPDFPortada(ciudad, idUnico)
                if portada["estado"]:
                    ruta_portada = portada["ruta"]
                    docs_eliminar.insert(0, ruta_portada)
                log_paquete = Hotel.generar_pdf_paquete(data["hotel"], data["actividades"] ,ticket, idUnico)
                if log_paquete["estado"]:
                    ruta_paquete = log_paquete["ruta"]
                    docs_eliminar.append(ruta_paquete)
                log_hotel = Hotel.generar_pdf_hotel(data["hotel"], idUnico)
                if log_hotel["estado"]:
                    incluye.append("hospedaje")
                    ruta_hotel = log_hotel["ruta"]
                    docs_eliminar.append(ruta_hotel)
                if "actividades" in data and data["actividades"]:
                    log_actividades = Actividad.generarPdfActividades(data["actividades"], idUnico)
                    if log_actividades["estado"]:
                        incluye.append("actividades")
                        ruta_actividades = log_actividades["ruta"]
                        docs_eliminar.append(ruta_actividades)
                if "costos" in data and data["costos"]:
                    costos = Costos.generarPdfCostos(data["costos"], incluye, idUnico)
                    if costos["estado"]:
                        ruta_costos = costos["ruta"]
                        docs_eliminar.append(ruta_costos)

        else:
            # opcion3
            if "vuelo" in data and data["vuelo"]:
                personas = {"pasajeros": data["vuelo"]["personas"].replace(",", "\n")}
                ticket = "Si incluye ticket(s) aéreo(s)"
                ciudad = "\n".join(item[f"ciudad_destino{index}"].split(",")[0] for index, item in enumerate(data["vuelo"]["segmentos"]) if index < len(data["vuelo"]["segmentos"]) - 1)
                portada = Cotizador.generarPDFPortada(ciudad, idUnico)
                if portada["estado"]:
                    ruta_portada = portada["ruta"]
                    docs_eliminar.append(ruta_portada)
                log_paquete = Hotel.generar_pdf_paquete(data["hotel"], data["actividades"], ticket,idUnico, ciudad, personas)
                if log_paquete["estado"]:
                    ruta_paquete = log_paquete["ruta"]
                    docs_eliminar.append(ruta_paquete)
                vuelos = Cotizador.cotizar_vuelos(data["vuelo"], idUnico)
                if vuelos["estado"]:
                    incluye.append("tickets aéreos")
                    ruta_vuelos = vuelos["ruta"]
                    docs_eliminar.append(ruta_vuelos)
                if "actividades" in data and data["actividades"]:
                    log_actividades = Actividad.generarPdfActividades(data["actividades"], idUnico)
                    if log_actividades["estado"]:
                        incluye.append("actividades")
                        ruta_actividades = log_actividades["ruta"]
                        docs_eliminar.append(ruta_actividades)
                if "costos" in data and data["costos"]:
                    costos = Costos.generarPdfCostos(data["costos"], incluye, idUnico)
                    if costos["estado"]:
                        ruta_costos = costos["ruta"]
                        docs_eliminar.append(ruta_costos)
            # opcion4
            else:
                return{"estado": False, "mensaje": "No hay datos de vuelos ni paquetes"}
        
        ruta_pdf = os.path.abspath(f"plantilla/cotizacion_completo_{idUnico}.pdf")
        log_unir_pdf = GenerarPdf.unir_pdfs(docs_eliminar, ruta_pdf)
        if log_unir_pdf:
            pdf_base64 = GenerarPdf.archivo_a_base64(ruta_pdf)
            if pdf_base64:
                docs_eliminar.append(ruta_pdf)
                log_eliminar_data = GenerarPdf.eliminar_documentos(docs_eliminar)
                if log_eliminar_data:
                    return {"estado": True, "mensaje": "Documento creado exitosamente", "pdf": pdf_base64}    
                else:
                    return {"estado": False, "mensaje": "No se logro crear base64"} 
            else:
                return {"estado": False, "mensaje": "No se logro eliminar los docs"} 
        else:
            return {"estado": False, "mensaje": "No se logro unir los docs generales"}
    


    
    
    @staticmethod
    def generarPDFPortada(ciudad, idUnico):
        if ciudad:
            datos={
                "city": ciudad.upper()
            }
            docs_eliminar = []
            ruta_plantilla_portada_original = os.path.abspath("plantilla/plantilla_cotizar_portada.docx")
            ruta_plantilla_portada = os.path.abspath(f"plantilla/plantilla_cotizar_portada_{idUnico}.docx")
            docs_eliminar.append(ruta_plantilla_portada)
            shutil.copy(ruta_plantilla_portada_original, ruta_plantilla_portada)
            ruta_docx_generado_portada = os.path.abspath(f"plantilla/portada_{idUnico}.docx")
            docs_eliminar.append(ruta_docx_generado_portada)
            longitud = len(ciudad)
            if longitud < 10:
                estilos = {"fuente": "Helvetica", "numero":60, "color": "#FFFFFF"}
            else:
                estilos = {"fuente": "Helvetica", "numero":40, "color": "#FFFFFF"}
            log_reemplazar_cotitazion = GenerarPdf.reemplazar_texto_docx(ruta_plantilla_portada, ruta_docx_generado_portada, datos, estilos, "CENTER")
            if log_reemplazar_cotitazion:
                ruta_directorio_pdf = os.path.abspath("plantilla")
                ruta_pdf_generado = GenerarPdf.convertir_docx_a_pdf(ruta_docx_generado_portada, ruta_directorio_pdf)
                if ruta_pdf_generado:
                    log_eliminar_data = GenerarPdf.eliminar_documentos(docs_eliminar)
                    if log_eliminar_data:
                        return{"estado": True, "mensaje": "Costos creados correctamente", "ruta": ruta_pdf_generado}
                else:
                    return {"estado": False, "mensaje": "No se ha podido convertir docx a pdf"} 
            else:
                return {"estado": False, "mensaje": "No se ha podido reemplazar docx"} 
        else:
            return {"estado": False, "mensaje": "No hay datos de ciudad"} 



    @staticmethod
    def cotizar_vuelos(data, idUnico):
        if data:
            docs_eliminar = []
            numero_segmentos = len(data["segmentos"])
            ruta_plantilla_cotizador_vuelos = os.path.abspath(f"plantilla/plantilla_cotizar_vuelos_{numero_segmentos}.docx")
            ruta_docx_generado_cotizacion_vuelos_segmento = os.path.abspath(f"plantilla/cotizacion_vuelos_seg_{idUnico}.docx")
            docs_eliminar.append(ruta_docx_generado_cotizacion_vuelos_segmento)
            shutil.copy(ruta_plantilla_cotizador_vuelos, ruta_docx_generado_cotizacion_vuelos_segmento)
            if numero_segmentos <= 2:
                estilos_tabla = {"fuente": "Helvetica", "numero":10}
            else:
                estilos_tabla = {"fuente": "Helvetica", "numero":7}
            aux = True
            resultado = {}
            for index, segmento in enumerate(data["segmentos"]):
                variable = (f"[detalle_vuelo{index}]")
                detalle_vuelo = segmento[f"detalle_vuelo{index}"]
                log_vuelos_legs = GenerarPdf.armar_tabla_vuelos(ruta_docx_generado_cotizacion_vuelos_segmento, ruta_docx_generado_cotizacion_vuelos_segmento,variable,detalle_vuelo ,estilos_tabla)
                if not log_vuelos_legs:
                    aux = False
                if variable in segmento:  # Verifica si la clave existe en el diccionario
                    segmento.pop(variable)
                resultado.update(segmento)
            if aux:
                data.pop("segmentos")
                resultado.update(data)
                ruta_aereolina_imagen = os.path.abspath(f"img/aereolinas_logos/{data['aereolina'].lower()}.png")
                if os.path.exists(ruta_aereolina_imagen):
                    GenerarPdf.imagen_en_docx(ruta_aereolina_imagen, ruta_docx_generado_cotizacion_vuelos_segmento, "[imagen_aereolina]", alto_en_pt=17)
                else:
                    resultado["imagen_aereolina"] = ""
                log_reemplazar_cotitazion = GenerarPdf.reemplazar_texto_docx(ruta_docx_generado_cotizacion_vuelos_segmento, ruta_docx_generado_cotizacion_vuelos_segmento, resultado, estilos_tabla)
                if log_reemplazar_cotitazion:
                    ruta_directorio_pdf = os.path.abspath("plantilla")
                    ruta_pdf_cotizacion_vuelos = GenerarPdf.convertir_docx_a_pdf(ruta_docx_generado_cotizacion_vuelos_segmento, ruta_directorio_pdf)
                    if ruta_pdf_cotizacion_vuelos:
                        log_eliminar_data = GenerarPdf.eliminar_documentos(docs_eliminar)
                        if log_eliminar_data:
                            return {"estado": True, "mensaje": "Documento creado exitosamente", "ruta": ruta_pdf_cotizacion_vuelos}    
                        else:
                            return {"estado": False, "mensaje": "No se logro eliminar los documentos auxiliares al crear los vuelos"}
                    else:
                        return {"estado": False, "mensaje": "No se ha podido convertir docx a pdf"} 
                else:
                    return {"estado": False, "mensaje": "No se ha posido reemplazar los datos en la plantilla"}
            else:
                return {"estado": False, "mensaje": "No se ha posido reemplazar las escalas en la plantilla"}   
        else:
            return {"estado": False, "mensaje": "No hay datos en el body"}  

class Adendum:
    @staticmethod
    def generar_adendum(data):
        if data:
            ruta_plantilla_adendum = os.path.abspath("plantilla/plantilla_adendum.docx")
            ruta_plantilla_declaraciones = os.path.abspath("plantilla/plantilla_declaraciones.docx")
            ruta_docx_generado_adendum = os.path.abspath("plantilla/adendum.docx")
            ruta_docx_generado_declaraciones = os.path.abspath("plantilla/declaraciones.docx")
            estilos = {"fuente": "Helvetica", "numero":10}
            log_reemplazar_adendum = GenerarPdf.reemplazar_texto_docx(ruta_plantilla_adendum, ruta_docx_generado_adendum, data, estilos)
            log_reemplazar_declaraciones = GenerarPdf.reemplazar_texto_docx(ruta_plantilla_declaraciones, ruta_docx_generado_declaraciones, data, estilos)
            if log_reemplazar_adendum and log_reemplazar_declaraciones:
                ruta_directorio_pdf = os.path.abspath("plantilla")
                ruta_pdf_adendum = GenerarPdf.convertir_docx_a_pdf(ruta_docx_generado_adendum, ruta_directorio_pdf)
                ruta_pdf_declaraciones = GenerarPdf.convertir_docx_a_pdf(ruta_docx_generado_declaraciones, ruta_directorio_pdf)
                if ruta_pdf_adendum and ruta_pdf_declaraciones:
                    log_imagenes = GenerarPdf.procesar_imagenes(data["recibos_pago"])
                    if log_imagenes["estado"] == True:
                        pdfs_unir = [ruta_pdf_adendum, ruta_pdf_declaraciones, log_imagenes["ruta"]]
                        ruta_pdf = os.path.abspath("plantilla/adendum_completo.pdf")
                        log_unir_pdf = GenerarPdf.unir_pdfs(pdfs_unir, ruta_pdf)
                        if log_unir_pdf:
                            docs_eliminar = [ruta_docx_generado_adendum, ruta_docx_generado_declaraciones, ruta_pdf_adendum, ruta_pdf_declaraciones, log_imagenes["ruta"]]
                            log_eliminar_data = GenerarPdf.eliminar_documentos(docs_eliminar)
                            if log_eliminar_data:
                                pdf_base64 = GenerarPdf.archivo_a_base64(ruta_pdf)
                                if pdf_base64:
                                    return {"estado": True, "mensaje": "Documento creado exitosamente", "pdf": pdf_base64}    
                                else:
                                    return {"estado": False, "mensaje": "No se logro crear base64"}    
                            else:
                                return {"estado": False, "mensaje": "No se logro eliminar los documentos auxiliares"}
                        else:
                            return {"estado": False, "mensaje": "No se ha podido unir Adendum con Declaraciones"}    
                    else:
                        return log_imagenes
                else:
                    return {"estado": False, "mensaje": "Documento PDF no se puede crear"}    
            else:
                    return {"estado": False, "mensaje": "No se ha posido reemplazar los datos en la plantilla"}   
        else:
            return {"estado": False, "mensaje": "No hay datos en el body"}   
           
class Contrato:
    @staticmethod
    def generar_contrato(data):
        if data:
            ruta_plantilla_contratos = os.path.abspath("plantilla/plantilla_contratos.docx")
            ruta_plantilla_declaraciones = os.path.abspath("plantilla/plantilla_declaraciones.docx")
            ruta_docx_generado_contratos = os.path.abspath("plantilla/contratos.docx")
            ruta_docx_generado_declaraciones = os.path.abspath("plantilla/declaraciones.docx")
            estilos = {"fuente": "Helvetica", "numero":10}
            log_reemplazar_contratos = GenerarPdf.reemplazar_texto_docx(ruta_plantilla_contratos, ruta_docx_generado_contratos, data, estilos)
            log_reemplazar_declaraciones = GenerarPdf.reemplazar_texto_docx(ruta_plantilla_declaraciones, ruta_docx_generado_declaraciones, data, estilos)
            if log_reemplazar_contratos and log_reemplazar_declaraciones:
                ruta_docx_generado_contratos_estilos = os.path.abspath("plantilla/contratos_estilo.docx")
                log_estilos = GenerarPdf.aplicar_estilos_especificos(ruta_docx_generado_contratos, ruta_docx_generado_contratos_estilos)
                if log_estilos:
                    ruta_directorio_pdf = os.path.abspath("plantilla")
                    ruta_pdf_contratos = GenerarPdf.convertir_docx_a_pdf(ruta_docx_generado_contratos_estilos, ruta_directorio_pdf)
                    ruta_pdf_declaraciones = GenerarPdf.convertir_docx_a_pdf(ruta_docx_generado_declaraciones, ruta_directorio_pdf)
                    if ruta_pdf_contratos and ruta_pdf_declaraciones:
                        log_imagenes = GenerarPdf.procesar_imagenes(data["recibos_pago"])
                        if log_imagenes["estado"] == True:
                            pdfs_unir = [ruta_pdf_contratos, ruta_pdf_declaraciones, log_imagenes["ruta"]]
                            ruta_pdf = os.path.abspath("plantilla/contrato_completo.pdf")
                            log_unir_pdf = GenerarPdf.unir_pdfs(pdfs_unir, ruta_pdf)
                            if log_unir_pdf:
                                docs_eliminar = [ruta_docx_generado_contratos, ruta_docx_generado_declaraciones, ruta_pdf_contratos, ruta_pdf_declaraciones, log_imagenes["ruta"], ruta_docx_generado_contratos_estilos]
                                log_eliminar_data = GenerarPdf.eliminar_documentos(docs_eliminar)
                                if log_eliminar_data:
                                    pdf_base64 = GenerarPdf.archivo_a_base64(ruta_pdf)
                                    if pdf_base64:
                                        return {"estado": True, "mensaje": "Documento creado exitosamente", "pdf": pdf_base64}    
                                    else:
                                        return {"estado": False, "mensaje": "No se logro crear base64"}    
                                else:
                                    return {"estado": False, "mensaje": "No se logro eliminar los documentos auxiliares"}
                            else:
                                return {"estado": False, "mensaje": "No se ha podido unir contratos con Declaraciones"}    
                        else:
                            return log_imagenes
                    else:
                        return {"estado": False, "mensaje": "Documento PDF no se puede crear"}
                else:
                    return {"estado": False, "mensaje": "No se ha podido aplicar estilos"}    
            else:
                    return {"estado": False, "mensaje": "No se ha posido reemplazar los datos en la plantilla"}   
        else:
            return {"estado": False, "mensaje": "No hay datos en el body"}  

class GenerarPdf:
    @staticmethod
    def procesar_imagenes(imagenesBase64):
        if imagenesBase64:
            rutas_imagenes_pdf = []
            rutas_a_eliminar = []
            for indice, recibo in enumerate(imagenesBase64):
                ruta_imagen_aux = os.path.abspath("plantilla/recibo"+str(indice))
                ruta_imagen = GenerarPdf.guardar_imagen_base64(recibo, ruta_imagen_aux)
                rutas_a_eliminar.append(ruta_imagen)
                if ruta_imagen:
                    ruta_plantilla_imagen = os.path.abspath("plantilla/plantilla_imagenes.docx")
                    ruta_docx_imagen = os.path.abspath("plantilla/recibo"+str(indice)+".docx")
                    rutas_a_eliminar.append(ruta_docx_imagen)
                    log_imagen_docs = GenerarPdf.agregar_imagen_docx(ruta_imagen, ruta_plantilla_imagen,ruta_docx_imagen)
                    if log_imagen_docs:
                        ruta_directorio_pdf_imagen = os.path.abspath("plantilla")
                        ruta_imagen_pdf = GenerarPdf.convertir_docx_a_pdf(ruta_docx_imagen,ruta_directorio_pdf_imagen)
                        rutas_a_eliminar.append(ruta_imagen_pdf)
                        if ruta_imagen_pdf:
                            rutas_imagenes_pdf.append(ruta_imagen_pdf)
                        else:
                            return {"estado": False, "mensaje": "Hubo un error al tranformar imagenes a pdf"}  
                    else:
                        return {"estado": False, "mensaje": "Hubo un error al insertar las imagenes en el documento"}
                else:
                    return {"estado": False, "mensaje": "Hubo un error con las imagenes adjuntadas"}
            ruta_imagenes_unidas_pdf = os.path.abspath("plantilla/imagenes.pdf")
            print(rutas_imagenes_pdf)
            log_unir_imagenes = GenerarPdf.unir_pdfs(rutas_imagenes_pdf, ruta_imagenes_unidas_pdf)
            if log_unir_imagenes:
                GenerarPdf.eliminar_documentos(rutas_a_eliminar)
                return {"estado": True, "mensaje": "Se proceso bien las imagenes", "ruta":ruta_imagenes_unidas_pdf}
            else:
                return {"estado": False, "mensaje": "Hubo un error al unir las imagenes"}
        else:
            return {"estado": False, "mensaje": "No hay recibos de pago"}

    @staticmethod
    def archivo_a_base64(ruta_archivo):
        try:
            with open(ruta_archivo, "rb") as archivo:
                contenido = archivo.read()
                contenido_base64 = base64.b64encode(contenido).decode('utf-8')
                return contenido_base64
        except FileNotFoundError:
            logging.error(f"El archivo {ruta_archivo} no se encuentra.")
            print(f"El archivo {ruta_archivo} no se encuentra.")
            return False
        except Exception as e:
            logging.error(f"Error al convertir el PDF a base64: {e}")
            print(f"Error al convertir el PDF a base64: {e}")
            return False

    @staticmethod
    def unir_pdfs(rutas, ruta_resultado):
        try:
            merger = PdfMerger()
            for ruta in rutas:
                merger.append(ruta)
            merger.write(ruta_resultado)
            merger.close()
            return True
        except Exception as e:
            logging.error(f"Error al combinar los PDFs: {e}")
            print(f"Error al combinar los PDFs: {e}")
            return False

    @staticmethod
    def eliminar_documentos(rutas_documentos):
        for ruta in rutas_documentos:
            try:
                # Verificar si el archivo existe
                if os.path.exists(ruta):
                    os.remove(ruta)  # Eliminar el archivo
                else:
                    logging.error(f"El archivo {ruta} no existe.")
                    print(f"El archivo {ruta} no existe.")
            except Exception as e:
                logging.error(f"Error al intentar eliminar el archivo {ruta}: {e}")
                print(f"Error al intentar eliminar el archivo {ruta}: {e}")
                return False
        return True

    @staticmethod
    def reemplazar_texto_docx(archivo_entrada, archivo_salida, variables, estilos, alineacion="JUSTIFY"):
        try:
            doc = Document(archivo_entrada)
            for para in doc.paragraphs:
                for var, valor in variables.items():
                    if isinstance(valor, list):  # Si el valor es una lista (como paquete_incluye)
                        valor = "\n".join(str(item) if isinstance(item, dict) else item for item in valor)
                    marcador = f"[{var}]"
                    if marcador in para.text:
                        para.text = para.text.replace(marcador, str(valor))
                        for run in para.runs:
                            run.font.name = estilos["fuente"]
                            run.font.size = Pt(estilos["numero"])
                            if "color" in estilos and estilos["color"]:
                                color_hex = estilos["color"]
                                if color_hex:  # Si hay un color especificado
                                    # Convertir el color hexadecimal a RGB
                                    rgb = RGBColor(int(color_hex[1:3], 16), int(color_hex[3:5], 16), int(color_hex[5:7], 16))
                                    run.font.color.rgb = rgb
                        # Justificar el párrafo
                        if alineacion == "CENTER":
                            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        else:
                            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        

            #Recorrer las tablas del documento
            for tabla in doc.tables:
                for fila in tabla.rows:
                    for celda in fila.cells:
                        for para in celda.paragraphs:
                            for var, valor in variables.items():
                                if isinstance(valor, list):  # Si el valor es una lista
                                    valor = "\n".join(str(item) if isinstance(item, dict) else item for item in valor)
                                marcador = f"[{var}]"
                                if marcador in para.text:
                                    para.text = para.text.replace(marcador, str(valor))
                                    for run in para.runs:
                                        run.font.name = estilos["fuente"]
                                        run.font.size = Pt(estilos["numero"])
                                        if "color" in estilos and estilos["color"]:
                                            color_hex = estilos["color"]
                                            if color_hex:  # Si hay un color especificado
                                                # Convertir el color hexadecimal a RGB
                                                rgb = RGBColor(int(color_hex[1:3], 16), int(color_hex[3:5], 16), int(color_hex[5:7], 16))
                                                run.font.color.rgb = rgb
            doc.save(archivo_salida)
            return True
        except Exception as e:
            logging.error(f"Error: {e}")
            print(f"Error: {e}")  # Imprime el error si ocurre
            return False  # En caso de error, devolver False
        
        
    def reemplazar_clave_array(archivo_entrada, archivo_salida, array, estilos, clave, alineacion="K"):
        try:
            doc = Document(archivo_entrada)

            # Construir el texto a insertar a partir del array de strings
            texto_reemplazo = "\n".join(array)

            def aplicar_reemplazo(parrafo):
                """Reemplaza el texto y aplica estilo y alineación."""
                parrafo.text = texto_reemplazo
                for run in parrafo.runs:
                    run.font.name = estilos["fuente"]
                    run.font.size = Pt(estilos["numero"])
                if alineacion == "CENTER":
                    parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif alineacion == "JUSTIFY":
                    parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # Reemplazar en párrafos
            for para in doc.paragraphs:
                if clave in para.text:
                    aplicar_reemplazo(para)
            
            # Reemplazar en tablas
            for tabla in doc.tables:
                for fila in tabla.rows:
                    for celda in fila.cells:
                        for para in celda.paragraphs:
                            if clave in para.text:
                                aplicar_reemplazo(para)
            
            # Guardar el documento actualizado
            doc.save(archivo_salida)
            return True

        except Exception as e:
            logging.error(f"Error: {e}")
            print(f"Error: {e}")  # Imprime el error si ocurre
            return False  # En caso de error, devolver False
    

        
    @staticmethod
    def crear_tabla_rooms(archivo_entrada, archivo_salida, variable, datos, estilos):
        try:
            doc = Document(archivo_entrada)

            for para in doc.paragraphs:
                if variable in para.text:
                    para.clear()  # Eliminar el contenido del párrafo con la variable

                    for diccionario in datos:
                        # Crear una tabla para cada diccionario
                        table = doc.add_table(rows=5, cols=4)
                        table.style = 'Table Grid'
                        for row in table.rows:
                            for cell in row.cells:
                                tc = cell._element
                                tc_pr = tc.get_or_add_tcPr()
                                tc_borders = parse_xml(r'''
                                    <w:tcBorders %s>
                                        <w:top w:val="single" w:sz="4" w:space="0" w:color="E7E6E6"/>
                                        <w:left w:val="single" w:sz="4" w:space="0" w:color="E7E6E6"/>
                                        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E7E6E6"/>
                                        <w:right w:val="single" w:sz="4" w:space="0" w:color="E7E6E6"/>
                                    </w:tcBorders>''' % nsdecls('w'))
                                tc_pr.append(tc_borders)
                                # Establecer el color de fondo de la celda
                                shading_elm = parse_xml(r'<w:shd {} w:fill="E7E6E6"/>'.format(nsdecls('w')))
                                tc_pr.append(shading_elm)
                        comentarios = []
                        pares = list(diccionario.items())
                        for row_idx, row in enumerate(table.rows):
                            for col_idx, cell in enumerate(row.cells):
                                # Asignar las claves y valores según la posición
                                index = row_idx * 2 + col_idx // 2
                                if index < len(pares):
                                    clave, valor = pares[index]
                                    if clave == "rate_comments":
                                        comentarios.append(str(valor))
                                    else:
                                        if col_idx % 2 == 0:
                                            cell.text = GenerarPdf.traducir_palabras(clave)  # Clave en columna izquierda
                                            # Aplicar negrita en la columna de claves
                                            for paragraph in cell.paragraphs:
                                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                                                for run in paragraph.runs:
                                                    run.bold = True
                                        else:
                                            cell.text = str(valor)  # Valor en columna derecha

                                        # Aplicar estilos generales
                                        for paragraph in cell.paragraphs:
                                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                                            for run in paragraph.runs:
                                                run.font.name = estilos["fuente"]
                                                run.font.size = Pt(estilos["numero"])

                        # Insertar un párrafo vacío y luego la tabla
                        empty_paragraph = para.insert_paragraph_before()
                        empty_paragraph.text = ""  # Salto de línea entre tablas
                        table_element = table._element
                        empty_paragraph._element.addnext(table_element)  # Agregar la tabla después del párrafo vacío
                        if comentarios:
                            comentario_paragraph = empty_paragraph.insert_paragraph_before()
                            comentario_paragraph.text = "Comentarios: " + "; ".join(comentarios)
                            for run in comentario_paragraph.runs:
                                run.font.name = estilos["fuente"]
                                run.font.size = Pt(estilos["numero"])
                            table_element.addnext(comentario_paragraph._element)
            doc.save(archivo_salida)
            return True
        except Exception as e:
            logging.error(f"Error: {e}")
            print(f"Error: {e}")  # Imprime el error si ocurre
            return False  # En caso de error, devolver False

    @staticmethod
    def imagen_en_docx(image_path, docx_path, key, ancho_en_pt=None, alto_en_pt=None):
        try:
            doc = Document(docx_path)

            image = Image.open(image_path)
            
            
            # if image.mode == 'RGBA':
            #     # Crear una nueva imagen con fondo blanco
            #     fondo_blanco = Image.new("RGBA", image.size, (255, 255, 255, 255))  # Fondo blanco (RGBA)
            #     fondo_blanco.paste(image, (0, 0), image)  # Pegar la imagen original sobre el fondo blanco
            #     image = fondo_blanco.convert('RGB')

            # Calcular el ancho en proporción al alto especificado
            ancho_original, alto_original = image.size
            
            if ancho_en_pt is not None and alto_en_pt is None:
                # Si solo se proporciona el ancho, calcular el alto manteniendo la proporción
                proporción = ancho_en_pt / ancho_original
                alto_en_pt = int(alto_original * proporción)
            elif alto_en_pt is not None and ancho_en_pt is None:
                # Si solo se proporciona el alto, calcular el ancho manteniendo la proporción
                proporción = alto_en_pt / alto_original
                ancho_en_pt = int(ancho_original * proporción)
            else:
                return False


            extension = os.path.splitext(image_path)[1].lower()
            if extension == '.png':
                formato_imagen = 'PNG'
            else:
                formato_imagen = 'JPEG'

            # Guardar la imagen en un buffer de memoria
            image_stream = BytesIO()
            image.save(image_stream, format=formato_imagen)  # Guardar como JPEG
            image_stream.seek(0)

            # Bandera para verificar si se encontró la clave
            found = False

            # Buscar la clave en las celdas de la tabla
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if key in cell.text:
                            found = True
                            
                            # Limpiar el contenido de la celda
                            for paragraph in cell.paragraphs:
                                paragraph.clear()  # Limpiar el contenido del párrafo

                            # Insertar la imagen en la celda y centrarla
                            run = cell.paragraphs[0].add_run()
                            run.add_picture(image_stream, width=Pt(ancho_en_pt), height=Pt(alto_en_pt))
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar el párrafo
                            break
            
            # Buscar la clave en los párrafos normales si no está en una tabla
            if not found:
                for paragraph in doc.paragraphs:
                    if key in paragraph.text:
                        # Limpiar el contenido del párrafo
                        paragraph.clear()  # Limpiar el contenido del párrafo

                        # Crear un nuevo párrafo para la imagen
                        new_paragraph = paragraph.insert_paragraph_before()
                        new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar el párrafo
                        run = new_paragraph.add_run()
                        run.add_picture(image_stream, width=Pt(ancho_en_pt), height=Pt(alto_en_pt))
                        break

            # Guardar el documento actualizado
            doc.save(docx_path)
            return True

        except Exception as e:
            print(f"Error: {e}")
            traceback.print_exc()  # Imprime la traza del error
            return False

    @staticmethod
    def traducir_palabras(palabra):
        if palabra == "room_name":
            return "Nombre habitación:"
        elif palabra == "acomodation":
            return "Acomodation Type:"
        elif palabra == "name_pax":
            return "Pasajero:"
        elif palabra == "adults":
            return "Adultos:"
        elif palabra == "children":
            return "Niños"
        elif palabra == "age_children":
            return "Edad niños:"
        elif palabra == "board_basis":
            return "Board Base:"
        elif palabra == "room_number":
            return "Número habitaciones:"
        elif palabra == "rate_comments":
            return "Comentario de tarifa:"
        elif palabra == "Forma de pago":
            return "Comentario de tarifa:"
        else:
            return palabra
        


    def eliminar_filas_docx(ruta_entrada, ruta_salida, filas_a_eliminar):
        try:
            doc = Document(ruta_entrada)
            tabla = doc.tables[1]
            num_filas = len(tabla.rows)
            filas_a_eliminar_validas = [fila_idx for fila_idx in filas_a_eliminar if fila_idx < num_filas]
            for fila_idx in sorted(filas_a_eliminar_validas, reverse=True):
                fila = tabla.rows[fila_idx]
                tbl = tabla._element
                tbl.remove(fila._element)  # Elimina la fila del XML de la tabla
            
            doc.save(ruta_salida)
            return True
        
        except Exception as e:
            print(f"Error: {e}")
            return False


    @staticmethod
    def armar_tabla_vuelos(archivo_entrada, archivo_salida, variable,datos ,estilos):
        columnas = ["clase", "detalle_salida", "duracion", "detalle_destino"]
        numeroFilas = len(datos)
        ancho_columnas = [Inches(0.9), Inches(2.2), Inches(0.5), Inches(2.3)]
        try:
            doc = Document(archivo_entrada)
            for para in doc.paragraphs:
                if variable in para.text:
                    para.clear()
                    table = doc.add_table(rows=numeroFilas, cols=4)
                    table.style = 'Plain Table 2'

                    # Aplicar ancho a las columnas
                    for i, column in enumerate(table.columns):
                        column.width = ancho_columnas[i]
                    
                    # Primero agregamos las filas de datos a la tabla
                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            if i < len(datos):  # Validar que exista la fila
                                clave = columnas[j]  # Obtener la clave correspondiente a la columna
                                if clave in datos[i]:  # Validar que la clave exista en el diccionario
                                    valor = datos[i][clave]
                                    if isinstance(valor, list):
                                        valor = "\n".join(valor)
                                    cell.text = valor
                                    for paragraph in cell.paragraphs:
                                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    for run in paragraph.runs:
                                        run.bold = False  # Esto quita la negrita

                       # Verificar si escala tiene valor y agregar fila adicional
                        escala = datos[i].get("escala", "")
                        if escala:  # Si escala no está vacío
                            new_row = table.add_row()  # Agregar una nueva fila
                            new_row._tr.addnext(row._tr)  # Insertar después de la fila actual
                            merged_cell = new_row.cells[0]  # Seleccionar la primera celda de la fila
                            merged_cell.merge(new_row.cells[-1])  # Unir todas las columnas en una sola
                            merged_cell.text = escala  # Colocar el texto de escala
                            for paragraph in merged_cell.paragraphs:
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            for run in paragraph.runs:
                                run.bold = False

                    table_element = table._element
                    para._element.addnext(table_element)
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = estilos["fuente"]
                                    run.font.size = Pt(estilos["numero"])
            doc.save(archivo_salida)
            return True
        except Exception as e:
            logging.error(f"Error: {e}")
            print(f"Error: {e}")  # Imprime el error si ocurre
            return False  # En caso de error, devolver False

 
    @staticmethod
    def aplicar_estilos_especificos(archivo_entrada, archivo_salida):
        textos_objetivo = {
            'COMPAÑÍA TURISTICA "MARKETING VIP S.A" COMTUMARK',
            'EL CLIENTE',
            'SEGUNDA',
            'TERCERA'
            'LA OPERADORA TURISTICA',
            'Lugar y fecha'
        }
        try:
            doc = Document(archivo_entrada)
            for para in doc.paragraphs:
                nuevo_runs = []
                for run in para.runs:
                    texto = run.text
                    i = 0
                    while i < len(texto):
                        texto_encontrado = False
                        for palabra in textos_objetivo:
                            if texto[i:i+len(palabra)] == palabra:
                                # Crear un nuevo run con negrita para el texto objetivo
                                nuevo_run = para.add_run(palabra)
                                nuevo_run.bold = True
                                nuevo_run.font.name = 'Helvetica'
                                nuevo_run.font.size = Pt(10)
                                nuevo_runs.append(nuevo_run)
                                i += len(palabra)
                                texto_encontrado = True
                                break
                        if not texto_encontrado:
                            # Crear un run normal para el texto que no coincide
                            nuevo_run = para.add_run(texto[i])
                            nuevo_run.font.name = run.font.name or 'Helvetica'
                            nuevo_run.font.size = run.font.size or Pt(10)
                            nuevo_run.bold = run.bold
                            nuevo_runs.append(nuevo_run)
                            i += 1
                    run._element.getparent().remove(run._element)
            doc.save(archivo_salida)
            return True
        except Exception as e:
            logging.error(f"Error: {e}")
            print(f"Error: {e}")  # Imprime el error si ocurre
            return False  # En caso de error, devolver False

    @staticmethod
    def convertir_docx_a_pdf(archivo_entrada, archivo_salida):
        try:
            if not os.path.exists(archivo_entrada):
                raise FileNotFoundError(f"El archivo {archivo_entrada} no se encuentra.")
            if sys.platform == "win32":  # Windows
                libreoffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
            elif sys.platform == "linux" or sys.platform == "linux2":  # Linux
                libreoffice_path = "/usr/bin/libreoffice"  # Asegúrate de que esté en esa ubicación
            else:
                raise EnvironmentError("Sistema operativo no soportado")
            subprocess.run([libreoffice_path, '--headless', '--convert-to', 'pdf', archivo_entrada, '--outdir', archivo_salida])
            nombre_pdf = os.path.splitext(os.path.basename(archivo_entrada))[0] + ".pdf"
            ruta_pdf_salida = os.path.join(archivo_salida, nombre_pdf)
            if not os.path.exists(ruta_pdf_salida):
                logging.error(f"No se pudo crear el archivo PDF en {ruta_pdf_salida}.")
                return False
            return ruta_pdf_salida
        except Exception as e:
            logging.error(f"Error: {e}")
            print(f"Error: {e}")  # Imprime el error si ocurre
            return False  # En caso de error, devolver False

    @staticmethod
    def guardar_imagen_base64(imagen_base64, ruta_salida):
        try:
            match = re.match(r"data:image/(\w+);base64,(.*)", imagen_base64)
            if not match:
                return False
            formato_imagen = match.group(1)
            imagen_base64_data = match.group(2)
            imagen_bytes = base64.b64decode(imagen_base64_data)
            ruta_imagen = f"{ruta_salida}.{formato_imagen}"
            with open(f"{ruta_salida}.{formato_imagen}", "wb") as f:
                f.write(imagen_bytes)
            return ruta_imagen
        except Exception as e:
            logging.error(f"Error: {e}")
            print(f"Error al guardar la imagen: {e}")
            return False  

    @staticmethod
    def agregar_imagen_docx(ruta_imagen, ruta_plantilla, ruta_salida):
        try:
            # Cargar la plantilla del documento
            doc = Document(ruta_plantilla)
            
            # Abrir la imagen y obtener sus dimensiones
            with Image.open(ruta_imagen) as img:
                width, height = img.size
                
                # Ajustar dimensiones de la imagen
                if height > width:  # Si la imagen es más alta que ancha
                    alto = Cm(19)
                    ancho = (width / height) * alto
                else:  # Si la imagen es más ancha que alta
                    ancho = Cm(13)
                    alto = (height / width) * ancho
                
                # Insertar la imagen al principio del documento
                paragraph = doc.add_paragraph()  # Crear un nuevo párrafo al inicio
                run = paragraph.add_run()  # Crear un "run" en el párrafo
                run.add_picture(ruta_imagen, width=ancho, height=alto)  # Insertar la imagen con tamaño ajustado
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar el párrafo
            
            # Guardar el documento con la imagen añadida al principio
            doc.save(ruta_salida)
            return True
        except Exception as e:
            logging.error(f"Error: {e}")
            print(f"Error: {e}")  # Imprime el error si ocurre
            return False  # En caso de error, devolver False
        
    def download_image(image_url: str, save_path: str):
        try:
            # Realizar la solicitud HTTP para obtener la imagen
            response = requests.get(image_url)
            
            # Verificar que la solicitud fue exitosa
            response.raise_for_status()
            
            # Abrir la imagen desde los datos descargados
            image = Image.open(BytesIO(response.content))
            
            # Guardar la imagen en la ruta especificada sin modificar su tamaño
            image.save(save_path)
            
            return True
        
        except requests.exceptions.RequestException as e:
            # Manejar errores de la solicitud HTTP
            print(f"Error al descargar la imagen: {e}")
            return False
        except Exception as e:
            # Manejar otros posibles errores
            print(f"Error: {e}")
            return False
        
class Descarga:
    @staticmethod
    def descargar_documento(ruta):
        base64 = GenerarPdf.archivo_a_base64(ruta)
        if base64:
            return {"estado": True, "mensaje": "Archivo obtenido con exito", "archivo": base64}   
        else:
            return {"estado": False, "mensaje": "No se ha podido codificar el archivo"}   
        
class Guardar:
    @staticmethod
    def guardar_archivo(ruta_guardar, base64_string):
        try:
            contenido_binario = base64.b64decode(base64_string)
            archivo_docx = BytesIO(contenido_binario)
            doc = Document(archivo_docx)
            directorio = os.path.dirname(ruta_guardar)
            if not os.path.exists(directorio):
                os.makedirs(directorio)
            doc.save(ruta_guardar)
            return True
        except Exception as e:
            logging.error(f"Error al guardar el archivo: {e}")
            print(f"Error al guardar el archivo: {e}")
            return False 

class Imagen:
    @staticmethod
    def cotizar_vuelos(data):
        if data:
            ruta_plantilla = os.path.abspath("plantilla/plantilla.jpg")
            ruta_imagen_generada = os.path.abspath("plantilla/vuelos.jpg")
            imagen_original = Image.open(ruta_plantilla)
            imagen_copia = imagen_original.copy()
            imagen_copia.save(ruta_imagen_generada)
            Imagen.colocar_texto_a_imagen(data["ida_fecha"],(170,110),ruta_imagen_generada,ruta_imagen_generada,15)
            imagen_pequena = Imagen.sacar_logo_aereolina(data["aereolina_codigo"])
            Imagen.colocar_imagen_pequena(imagen_pequena, (33,20), ruta_imagen_generada, ruta_imagen_generada,90,40)
            Imagen.colocar_texto_a_imagen(data["aereolina_nombre"],(120,30),ruta_imagen_generada,ruta_imagen_generada,15)
            Imagen.colocar_texto_a_imagen(data["vuelta_fecha"],(170,320),ruta_imagen_generada,ruta_imagen_generada,15)
            texto_ida = data["codigo_salida"] + "-" + data["codigo_destino"]
            Imagen.colocar_texto_a_imagen(texto_ida,(700,20),ruta_imagen_generada,ruta_imagen_generada,15)
            texto_vuelta = data["codigo_destino"] + "-" + data["codigo_salida"]
            Imagen.colocar_texto_a_imagen(texto_vuelta,(700,35),ruta_imagen_generada,ruta_imagen_generada,15)
            altura = 149
            for index, vuelo in enumerate(data["vuelos_ida"]):
                if index<3:
                    idVuelo = "I"+str(index+1)
                    Imagen.colocar_texto_a_imagen(idVuelo,(96,altura),ruta_imagen_generada,ruta_imagen_generada,13)
                    texto_horas = data["codigo_salida"]+": "+vuelo["hora_salida"]+" ---> "+data["codigo_destino"]+": "+vuelo["hora_llegada"]
                    Imagen.colocar_texto_a_imagen(texto_horas,(144,altura),ruta_imagen_generada,ruta_imagen_generada,15)
                    Imagen.colocar_texto_a_imagen(vuelo["duracion"],(400,altura),ruta_imagen_generada,ruta_imagen_generada,15)
                    escalas = str(vuelo["numero_escalas"])+ " Escala(s)"
                    Imagen.colocar_texto_a_imagen(escalas,(500,altura),ruta_imagen_generada,ruta_imagen_generada,15)
                    ruta_personal = Imagen.sacar_equipaje("personal",int(vuelo["equipaje_personal"]))
                    ruta_carry = Imagen.sacar_equipaje("carry",int(vuelo["equipaje_carry"]))
                    ruta_bodega = Imagen.sacar_equipaje("bodega",int(vuelo["equipaje_bodega"]))
                    Imagen.colocar_imagen_pequena(ruta_personal, (700,altura-3), ruta_imagen_generada, ruta_imagen_generada,18,18)
                    Imagen.colocar_imagen_pequena(ruta_carry, (720,altura-4), ruta_imagen_generada, ruta_imagen_generada,20,20)
                    Imagen.colocar_imagen_pequena(ruta_bodega, (740,altura-4), ruta_imagen_generada, ruta_imagen_generada,20,20)
                    altura = altura +52
            altura = 358
            for index, vuelo in enumerate(data["vuelos_vuelta"]):
                if index<3:
                    idVuelo = "v"+str(index+1)
                    Imagen.colocar_texto_a_imagen(idVuelo,(96,altura),ruta_imagen_generada,ruta_imagen_generada,13)
                    texto_horas = data["codigo_salida"]+": "+vuelo["hora_salida"]+" ---> "+data["codigo_destino"]+": "+vuelo["hora_llegada"]
                    Imagen.colocar_texto_a_imagen(texto_horas,(144,altura),ruta_imagen_generada,ruta_imagen_generada,15)
                    Imagen.colocar_texto_a_imagen(vuelo["duracion"],(400,altura),ruta_imagen_generada,ruta_imagen_generada,15)
                    escalas = str(vuelo["numero_escalas"])+ " Escala(s)"
                    Imagen.colocar_texto_a_imagen(escalas,(500,altura),ruta_imagen_generada,ruta_imagen_generada,15)
                    ruta_personal = Imagen.sacar_equipaje("personal",int(vuelo["equipaje_personal"]))
                    ruta_carry = Imagen.sacar_equipaje("carry",int(vuelo["equipaje_carry"]))
                    ruta_bodega = Imagen.sacar_equipaje("bodega",int(vuelo["equipaje_bodega"]))
                    Imagen.colocar_imagen_pequena(ruta_personal, (700,altura-3), ruta_imagen_generada, ruta_imagen_generada,18,18)
                    Imagen.colocar_imagen_pequena(ruta_carry, (720,altura-4), ruta_imagen_generada, ruta_imagen_generada,20,20)
                    Imagen.colocar_imagen_pequena(ruta_bodega, (740,altura-4), ruta_imagen_generada, ruta_imagen_generada,20,20)
                    altura = altura +52
            imagen_base64 = Imagen.convertir_imagen_a_base64(ruta_imagen_generada)
            if imagen_base64:
                return {"estado": True, "mensaje": "Imagen generada correctamente", "imagen": imagen_base64}  
            else:
                return {"estado": False, "mensaje": "No se ha podido generar Imagen"}  
        else:
            return {"estado": False, "mensaje": "No hay datos en el body"}  
        

    @staticmethod
    def convertir_imagen_a_base64(ruta_imagen):
        try:
            with open(ruta_imagen, "rb") as imagen:
                datos_imagen = imagen.read()
                base64_imagen = base64.b64encode(datos_imagen).decode("utf-8")
            return base64_imagen
        except Exception as e:
            print(f"Ocurrió un error: {e}")
            logging.error(f"Ocurrió un error: {e}")
            return False


    
    @staticmethod
    def colocar_texto_a_imagen(texto,coordenadas,ruta_imagen, ruta_salida,fuente):
        try:
            sistema_operativo = platform.system()
        
            # Configurar la ruta de la fuente según el sistema operativo
            if sistema_operativo == "Windows":
                ruta_fuente = "C:/Windows/Fonts/arial.ttf"  # Ruta típica en Windows
            elif sistema_operativo == "Linux":
                ruta_fuente = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"  # Fuente común en Linux
            else:
                logging.error("Sistema operativo no soportado para fuentes predeterminadas.")
                raise OSError("Sistema operativo no soportado para fuentes predeterminadas.")
            
            # Verificar que la fuente exista
            if not os.path.exists(ruta_fuente):
                logging.error(f"La fuente no se encuentra en {ruta_fuente}")
                raise FileNotFoundError(f"La fuente no se encuentra en {ruta_fuente}")

            imagen = Image.open(ruta_imagen)
            draw = ImageDraw.Draw(imagen)
            fuente = ImageFont.truetype(ruta_fuente, fuente)
            draw.text(coordenadas, texto, fill="black", font=fuente)
            imagen.save(ruta_salida)
            return True
        except Exception as e:
            print(f"Ocurrió un error: {e}") 
            logging.error(f"Ocurrió un error: al colocar texto{e}")
            return False



    @staticmethod
    def sacar_logo_aereolina(aereolina):
        if aereolina == "AV" or aereolina == '2K':
            return os.path.abspath("img/aereolinas_logos/avianca.png")
        elif aereolina == 'CM':
            return os.path.abspath("img/aereolinas_logos/copa.png")
        elif aereolina == 'DL':
            return os.path.abspath("img/aereolinas_logos/delta.png")        
        elif aereolina == 'B6':
            return os.path.abspath("img/aereolinas_logos/jet.png")
        elif aereolina == 'LA':
            return os.path.abspath("img/aereolinas_logos/latam.png")
        elif aereolina == 'AA':
            return os.path.abspath("img/aereolinas_logos/american.png")
        

    @staticmethod
    def sacar_equipaje(tipo,id):
        if tipo == "personal":
            if id >= 1:
                return os.path.abspath("img/equipaje/si_personal.png")
            else:
                return os.path.abspath("img/equipaje/no_personal.png")
        elif tipo == 'carry':
            if id >= 1:
                return os.path.abspath("img/equipaje/si_carry.png")
            else:
                return os.path.abspath("img/equipaje/no_carry.png")
        elif tipo == 'bodega':
            if id >= 1:
                return os.path.abspath("img/equipaje/si_bodega.png")
            else:
                return os.path.abspath("img/equipaje/no_bodega.png")       
        


    @staticmethod
    def colocar_imagen_pequena(imagen_pequena, coordenadas, ruta_imagen, ruta_salida, ancho_pequena, alto_pequena):
        try:
            # Cargar la imagen de fondo
            imagen_grande = Image.open(ruta_imagen)
            
            # Crear una nueva imagen blanca del mismo tamaño que la imagen de fondo
            fondo_blanco = Image.new("RGB", imagen_grande.size, (255, 255, 255))
            
            # Pegar la imagen grande en el fondo blanco
            fondo_blanco.paste(imagen_grande, (0, 0))
            
            # Cargar la imagen pequeña
            imagen_pequena = Image.open(imagen_pequena)
            
            # Redimensionar la imagen pequeña al tamaño especificado
            imagen_pequena = imagen_pequena.resize((ancho_pequena, alto_pequena), Image.LANCZOS)  # Cambiado de ANTIALIAS a LANCZOS
            
            # Pegar la imagen pequeña en el fondo blanco en las coordenadas deseadas
            fondo_blanco.paste(imagen_pequena, coordenadas, 
                            imagen_pequena.convert("RGBA").getchannel("A") if imagen_pequena.mode == 'RGBA' else None)
            
            # Guardar la imagen resultante
            fondo_blanco.save(ruta_salida)            
        except Exception as e:
            print(f"Ocurrió un error: {e}") 
            logging.error(f"Ocurrió un error al colocar imagenes: {e}")  

class Hotel:
    @staticmethod
    def generar_voucher(data):
        if data:
            rooms = data["rooms"]
            data.pop("rooms")
            ruta_plantilla_voucher = os.path.abspath("plantilla/plantilla_voucher_hotel.docx")
            ruta_docx_generado_tabla = os.path.abspath("plantilla/voucher_tabla.docx")
            estilos = {"fuente": "Helvetica", "numero":10}
            log_tabla_rooms = GenerarPdf.crear_tabla_rooms(ruta_plantilla_voucher,ruta_docx_generado_tabla,"[rooms]", rooms, estilos)
            if log_tabla_rooms:
                ruta_docx_generado_voucher = os.path.abspath("plantilla/voucher.docx")
                log_reemplazar_cotitazion = GenerarPdf.reemplazar_texto_docx(ruta_docx_generado_tabla, ruta_docx_generado_voucher, data, estilos)
                if log_reemplazar_cotitazion:
                    ruta_directorio_pdf = os.path.abspath("plantilla")
                    ruta_pdf_cotizacion_vuelos = GenerarPdf.convertir_docx_a_pdf(ruta_docx_generado_voucher, ruta_directorio_pdf)
                    if ruta_pdf_cotizacion_vuelos:
                        docs_eliminar = [ruta_docx_generado_voucher,ruta_docx_generado_tabla]
                        log_eliminar_data = GenerarPdf.eliminar_documentos(docs_eliminar)
                        if log_eliminar_data:
                            pdf_base64 = GenerarPdf.archivo_a_base64(ruta_pdf_cotizacion_vuelos)
                            if pdf_base64:
                                return {"estado": True, "mensaje": "Documento creado exitosamente", "pdf": pdf_base64}    
                            else:
                                return {"estado": False, "mensaje": "No se logro crear base64"}    
                        else:
                            return {"estado": False, "mensaje": "No se logro eliminar los documentos auxiliares"}
                    else:
                        return {"estado": False, "mensaje": "No se ha podido convertir docx a pdf"}   
                else:
                    return {"estado": False, "mensaje": "No se ha posido reemplazar los datos en la plantilla"}
            else:
                return {"estado": False, "mensaje": "No se logro armar la tabla"} 
        else:
            return {"estado": False, "mensaje": "No hay datos en el body"}


    
    @staticmethod
    def generar_pdf_paquete(dataHotel, actividades, ticket, idUnico, ciudad=None, personas = None):
        docs_eliminar = []
        if dataHotel:
            numero_hoteles = len(dataHotel)
            if numero_hoteles <= 2:
                estilos = {"fuente": "Helvetica", "numero":12}
            elif numero_hoteles == 3:
                estilos = {"fuente": "Helvetica", "numero":9}
            elif numero_hoteles == 4:
                estilos = {"fuente": "Helvetica", "numero":7}
            else:
                return {"estado": False, "mensaje": "Ha excedido el numero de hoteles"}
            ruta_plantilla_paquete = os.path.abspath(f"plantilla/plantilla_cotizar_paquete_{numero_hoteles}.docx")
            ruta_docx_generado_paquete = os.path.abspath(f"plantilla/cotizacion_paquete_{idUnico}.docx")
            shutil.copy(ruta_plantilla_paquete, ruta_docx_generado_paquete)
            docs_eliminar.append(ruta_docx_generado_paquete)
            resultado = {}
            habitaciones = {}
            act = {}
            for index, hotel in enumerate(dataHotel):
                habitaciones[f"habitacion{index}"] = []
                act[f"actividades{index}"] = []
                adultos = 0
                ninos = 0
                dias = Hotel.calcular_dias_noches(hotel["check_in"], hotel["check_out"])
                for room in hotel["rooms"]:
                    adultos = adultos + (int(room["adults"])*int(room["room_number"]))
                    ninos = ninos + (int(room["children"])*int(room["room_number"]))
                    detalle_habitacion = (f"{room['room_number']} habitacion(es) {room['acomodation']} {room['room_name']} - {room['board_basis']}")
                    habitaciones[f"habitacion{index}"].append(detalle_habitacion)
                for item in actividades:
                    if item["codigo"] == hotel["city_code"]:
                        for tour in item["tours"]:
                            act[f"actividades{index}"].append(tour["nombre"])
                        break  # Romper el ciclo si ya se encuentra el id
                if not act[f"actividades{index}"]:  # Verifica si está vacío
                    act[f"actividades{index}"].append("No incluye actividades") 
                pax = "" 
                if adultos>1: pax += (f"{adultos} adulto(s)")
                if ninos>1: pax += (f"\n{ninos} niños(s)")
                datos = {
                    f"pasajeros": pax,
                    f"city{index}": hotel["city"],
                    f"dias{index}": (f"{dias['dias']} dias y {dias['noches']} noches"),
                    f"check_in{index}": hotel["check_in"],
                    f"check_out{index}": hotel["check_out"],
                    f"ticket{index}": ticket
                }
                resultado.update(datos)
                GenerarPdf.reemplazar_clave_array(ruta_docx_generado_paquete, ruta_docx_generado_paquete, act[f"actividades{index}"], estilos, f"[actividades{index}]")
                GenerarPdf.reemplazar_clave_array(ruta_docx_generado_paquete, ruta_docx_generado_paquete, habitaciones[f"habitacion{index}"], estilos, f"[habitacion{index}]")
        else:
            ruta_plantilla_paquete = os.path.abspath(f"plantilla/plantilla_cotizar_paquete_0.docx")
            ruta_docx_generado_paquete = os.path.abspath(f"plantilla/cotizacion_paquete_{idUnico}.docx")
            shutil.copy(ruta_plantilla_paquete, ruta_docx_generado_paquete)
            estilos = {"fuente": "Helvetica", "numero":12}
            docs_eliminar.append(ruta_docx_generado_paquete)
            act = ""
            if actividades:
                act = "\n".join(f"{tour['nombre']} - {item['ciudad']}" for item in actividades for tour in item["tours"])
            else:
                act = "No incluye actividades" 
            resultado = {
                "city": ciudad,
                "ticket": ticket,
                "actividades": act
            }
            resultado.update(personas)
        log_reemplazar_paquete = GenerarPdf.reemplazar_texto_docx(ruta_docx_generado_paquete,ruta_docx_generado_paquete, resultado, estilos)
        if log_reemplazar_paquete:
            ruta_directorio_pdf = os.path.abspath("plantilla")
            ruta_pdf_cotizacion_paquete = GenerarPdf.convertir_docx_a_pdf(ruta_docx_generado_paquete, ruta_directorio_pdf)
            if ruta_pdf_cotizacion_paquete:
                log_eliminar_data = GenerarPdf.eliminar_documentos(docs_eliminar)
                if log_eliminar_data:
                    return {"estado": True, "mensaje": "Documento creado exitosamente", "ruta": ruta_pdf_cotizacion_paquete} 
            else:
                return {"estado": False, "mensaje": "No se ha podido convertir docx a pdf"}                     
        else:
            return {"estado": False, "mensaje": "No se ha podido reemplazar docx"}

            
        

    @staticmethod
    def generar_pdf_hotel(dataHotel, idUnico):
        if dataHotel:
            docs_eliminar = []
            docs_unir = []
            ruta_plantilla_hotel_original = os.path.abspath("plantilla/plantilla_cotizar_hoteles.docx")
            ruta_plantilla_hotel = os.path.abspath(f"plantilla/plantilla_cotizar_hoteles_{idUnico}.docx")
            shutil.copy(ruta_plantilla_hotel_original, ruta_plantilla_hotel)
            docs_eliminar.append(ruta_plantilla_hotel)
            estilos = {"fuente": "Helvetica", "numero":12}
            aux = True
            for index, hotel in enumerate(dataHotel):
                ruta_docx_generado_hotel = os.path.abspath(f"plantilla/cotizacion_hotel_{index}_{idUnico}.docx")
                facilidades = ", ".join(hotel["facilities"]) + "."
                datos = {
                    "nombre_hotel": hotel["hotel_name"],
                    "descripcion": hotel["descripcion"],
                    "ciudad_hotel": hotel["city"],
                    "facilities": facilidades
                }
                log_reemplazar_paquete = GenerarPdf.reemplazar_texto_docx(ruta_plantilla_hotel,ruta_docx_generado_hotel, datos, estilos)
                docs_eliminar.append(ruta_docx_generado_hotel)
                if log_reemplazar_paquete:
                    ruta_imagen_descargada = os.path.abspath(f"plantilla/imagen_hotel_{index}_{idUnico}.jpeg")
                    ruta_imagen = hotel["imagen"]
                    log_descargar_imagen = GenerarPdf.download_image(ruta_imagen, ruta_imagen_descargada)
                    if not log_descargar_imagen:
                        ruta_imagen_descargada = os.path.abspath(f"img/mkv.jpg")
                    else:
                        docs_eliminar.append(ruta_imagen_descargada)
                    log_imagen_docx = GenerarPdf.imagen_en_docx(ruta_imagen_descargada, ruta_docx_generado_hotel, "[imagen_hotel]", alto_en_pt=250)
                    if log_imagen_docx:
                        ruta_directorio_pdf = os.path.abspath("plantilla")
                        ruta_pdf_cotizacion_hotel = GenerarPdf.convertir_docx_a_pdf(ruta_docx_generado_hotel, ruta_directorio_pdf)
                        if ruta_pdf_cotizacion_hotel:
                            docs_eliminar.append(ruta_pdf_cotizacion_hotel)
                            docs_unir.append(ruta_pdf_cotizacion_hotel)
                    else:
                        aux = False
                        return {"estado": False, "mensaje": "No se ha podido añadir imagen a docx"}  
                else:
                    aux = False
                    return {"estado": False, "mensaje": "No se ha podido reemplazar texto en plantilla hotel"}
            if aux:
                ruta_hoteles_unidos = os.path.abspath(f"plantilla/hoteles_{idUnico}.pdf")
                log_unir_hoteles = GenerarPdf.unir_pdfs(docs_unir, ruta_hoteles_unidos)
                if log_unir_hoteles:
                    log_eliminar_data = GenerarPdf.eliminar_documentos(docs_eliminar)
                    if log_eliminar_data:
                        return {"estado": True, "mensaje": "Documento creado exitosamente", "ruta": ruta_hoteles_unidos}    
                else:
                    return {"estado": False, "mensaje": "No se logro unir los pdfs de hoteles"}    
            else:
                return {"estado": False, "mensaje": "No se ha podido convertir docx a pdf"}
        else:
            return {"estado": False, "mensaje": "No hay datos del hotel"}

        
    @staticmethod
    def calcular_dias_noches(check_in, check_out):
        # Convertir las fechas de string a objetos datetime
        formato = "%Y-%m-%d"  # Formato de la fecha (yyyy-mm-dd)
        fecha_in = datetime.strptime(check_in, formato)
        fecha_out = datetime.strptime(check_out, formato)
        
        # Incluir el día del check-out en el cálculo
        diferencia = (fecha_out - fecha_in).days + 1  # Sumar 1 para incluir el día de salida
        
        # Días y noches
        dias = diferencia
        noches = dias - 1  # Las noches son un día menos que los días de estancia
        
        # Devolver un diccionario con los resultados
        return {
            "dias": dias,
            "noches": noches
        }
        
class Actividad:
    @staticmethod
    def generarPdfActividades(dataActividades, idUnico):
        # return{"estado": True}
        if dataActividades:
            pdfs_unir = []
            docs_eliminar = []
            aux = True
            ruta_plantilla_actividad_original = os.path.abspath("plantilla/plantilla_cotizar_actividades.docx")
            ruta_plantilla_actividad = os.path.abspath(f"plantilla/plantilla_cotizar_actividades_{idUnico}.docx")
            shutil.copy(ruta_plantilla_actividad_original, ruta_plantilla_actividad)
            docs_eliminar.append(ruta_plantilla_actividad)
            for indice, actividad in enumerate(dataActividades):
                for index, act in enumerate(actividad["tours"]):
                    if act['nombre'] != "Transfer":
                        act["ciudad"] = actividad["ciudad"]
                        ruta_docx_generado_actividad = os.path.abspath(f"plantilla/actividad_{indice}_{index}_{idUnico}.docx")
                        docs_eliminar.append(ruta_docx_generado_actividad)
                        estilos = {"fuente": "Helvetica", "numero":12}
                        log_reemplazar_cotitazion = GenerarPdf.reemplazar_texto_docx(ruta_plantilla_actividad, ruta_docx_generado_actividad, act, estilos)
                        if log_reemplazar_cotitazion:
                            ruta_imagen_descargada = os.path.abspath(f"plantilla/imagen_actividad_{indice}_{index}_{idUnico}.jpeg")
                            ruta_imagen = (f"https://cotizador.mvevip.com/img/actividades_internas/{actividad['codigo']}/{act['id']}.jpg")
                            log_download = GenerarPdf.download_image(ruta_imagen, ruta_imagen_descargada)
                            if not log_download:
                                ruta_imagen_descargada = os.path.abspath(f"img/mkv.jpg")
                            else:
                                docs_eliminar.append(ruta_imagen_descargada)
                            GenerarPdf.imagen_en_docx(ruta_imagen_descargada, ruta_docx_generado_actividad, "[imagen_actividad]", ancho_en_pt=400)
                            ruta_directorio_pdf = os.path.abspath("plantilla")
                            ruta_pdf_generado = GenerarPdf.convertir_docx_a_pdf(ruta_docx_generado_actividad, ruta_directorio_pdf)
                            if ruta_pdf_generado:
                                pdfs_unir.append(ruta_pdf_generado)
                                docs_eliminar.append(ruta_pdf_generado)
                                aux = True
                            else:
                                aux = False 
                        else:
                            aux = False 
        if aux is True:
            ruta_pdf = os.path.abspath(f"plantilla/cotizar_actividades_{idUnico}.pdf")
            log_unir = GenerarPdf.unir_pdfs(pdfs_unir, ruta_pdf)
            if log_unir:
                log_eliminar_data = GenerarPdf.eliminar_documentos(docs_eliminar)
                if log_eliminar_data:
                    return {"estado": True, "mensaje": "Actividades creadas correctamente", "ruta": ruta_pdf}  
                else:
                    return {"estado": False, "mensaje": "No se logro eliminar los documentos auxiliares"}
            else:
                return {"estado": False, "mensaje": "No se ha podido convertir docx a pdf"} 
        else:
            return {"estado": False, "mensaje": "Hubo un error con las actividades"} 
        

class Costos:
    @staticmethod
    def generarPdfCostos(dataCostos, incluye, idUnico):
        if dataCostos:
            docs_eliminar = []
            filas = []
            incluye_paquete = ", ".join(incluye) + "."
            datos = {
                "fecha": datetime.now().strftime("%d de %B de %Y a las %H:%M horas"),
                "incluye": incluye_paquete,
            }
            if dataCostos["tipo"] == "0":
                ruta_plantilla_costos_original = os.path.abspath("plantilla/plantilla_cotizar_costos_detallado.docx")
                if dataCostos["detallado"]["adultos"]["numero"]>0:
                    data = {
                        "numA": dataCostos["detallado"]["adultos"]["numero"],
                        "precioAdulto_u": round(float(dataCostos["detallado"]["adultos"]["precio"]) / float(dataCostos["detallado"]["adultos"]["numero"]), 2) if float(dataCostos["detallado"]["adultos"]["numero"]) > 0 else 0,
                        "precioAdulto": round(float(dataCostos["detallado"]["adultos"]["precio"]), 2),
                    }
                    datos.update(data)
                else:
                    filas.append(1)
                if dataCostos["detallado"]["ninos"]["numero"]>0:
                    data = {
                        "numN": dataCostos["detallado"]["ninos"]["numero"],
                        "precioNino_u": round(float(dataCostos["detallado"]["ninos"]["precio"]) / float(dataCostos["detallado"]["ninos"]["numero"]), 2) if float(dataCostos["detallado"]["ninos"]["numero"]) > 0 else 0,
                        "precioNino": round(float(dataCostos["detallado"]["ninos"]["precio"]), 2),
                    }
                    datos.update(data)
                else:
                    filas.append(2)
                if dataCostos["detallado"]["infantes"]["numero"]>0:
                    data = {
                        "numI": dataCostos["detallado"]["infantes"]["numero"],
                        "precioInf_u": round(float(dataCostos["detallado"]["infantes"]["precio"]) / float(dataCostos["detallado"]["infantes"]["numero"]), 2) if float(dataCostos["detallado"]["infantes"]["numero"]) > 0 else 0,
                        "precioInf": round(float(dataCostos["detallado"]["infantes"]["precio"]), 2),
                    }
                    datos.update(data)
                else:
                    filas.append(3)
                if dataCostos["detallado"]["terceraEdad"]["numero"]>0:
                    data = {
                        "numT": dataCostos["detallado"]["terceraEdad"]["numero"],
                        "precioTer_u": round(float(dataCostos["detallado"]["terceraEdad"]["precio"]) / float(dataCostos["detallado"]["terceraEdad"]["numero"]), 2) if float(dataCostos["detallado"]["terceraEdad"]["numero"]) > 0 else 0,
                        "precioTer": round(float(dataCostos["detallado"]["terceraEdad"]["precio"]), 2),
                    }
                    datos.update(data)
                else:
                    filas.append(4)
                if dataCostos["detallado"]["discapacitados"]["numero"]>0:
                    data = {
                        "numD": dataCostos["detallado"]["discapacitados"]["numero"],
                        "precioDis_u": round(float(dataCostos["detallado"]["discapacitados"]["precio"]) / float(dataCostos["detallado"]["discapacitados"]["numero"]), 2) if float(dataCostos["detallado"]["discapacitados"]["numero"]) > 0 else 0,
                        "precioDis": round(float(dataCostos["detallado"]["discapacitados"]["precio"]), 2),
                    }
                    datos.update(data)
                else:
                    filas.append(5)
                data = {"total": round(dataCostos["detallado"]["total"] , 2)}
                datos.update(data)

            elif dataCostos["tipo"] == "1":
                ruta_plantilla_costos_original = os.path.abspath("plantilla/plantilla_cotizar_costos_no_detallado.docx")
                data = {
                    "paquete": round(float(dataCostos["noDetallado"]["paquete"]),2),
                    "vuelo": round(float(dataCostos["noDetallado"]["vuelo"]),2),
                    "total": round(float(dataCostos["noDetallado"]["paquete"]) + float(dataCostos["noDetallado"]["vuelo"]), 2)
                }
                datos.update(data)
            else:
                return {"estado": False, "mensaje": "No se logro identificar el tipo de costos"}
            
            ruta_plantilla_costos = os.path.abspath(f"plantilla/plantilla_cotizar_costos_{idUnico}.docx")
            shutil.copy(ruta_plantilla_costos_original, ruta_plantilla_costos)
            ruta_docx_generado_costos = os.path.abspath(f"plantilla/costos_{idUnico}.docx")
            docs_eliminar.append(ruta_plantilla_costos)
            docs_eliminar.append(ruta_docx_generado_costos)
            estilos = {"fuente": "Helvetica", "numero":12}
            if len(filas) > 0:
                log_tabla_emiminar_fila = GenerarPdf.eliminar_filas_docx(ruta_plantilla_costos,ruta_plantilla_costos,filas)
                if not log_tabla_emiminar_fila:
                    return {"estado": False, "mensaje": "No se logro eliminar las filas de costos"} 
            log_reemplazar_cotitazion = GenerarPdf.reemplazar_texto_docx(ruta_plantilla_costos, ruta_docx_generado_costos, datos, estilos)
            if log_reemplazar_cotitazion:
                ruta_directorio_pdf = os.path.abspath("plantilla")
                ruta_pdf_generado = GenerarPdf.convertir_docx_a_pdf(ruta_docx_generado_costos, ruta_directorio_pdf)
                if ruta_pdf_generado:
                    log_eliminar_data = GenerarPdf.eliminar_documentos(docs_eliminar)
                    if log_eliminar_data:
                        return{"estado": True, "mensaje": "Costos creados correctamente", "ruta": ruta_pdf_generado}
                else:
                    return {"estado": False, "mensaje": "No se ha podido convertir docx a pdf"} 
            else:
                return {"estado": False, "mensaje": "No se ha podido reemplazar docx"} 
        else:
            return {"estado": False, "mensaje": "No hay datos de costos"} 

                